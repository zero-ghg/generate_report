import json
import re
import subprocess
import sys
from base64 import b64decode
from copy import deepcopy
from datetime import datetime
from html import escape
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from urllib.parse import unquote

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shape import CT_Inline
from docx.shared import Cm, Pt, RGBColor


FONT_NAME = "宋体"
EMPTY = "—"
COMPACT_TABLE_FONT_SIZE = 5.5
COVER_UNDERLINE_TAB_POS = "7200"
SUBPROJECT_SECTION_HEADING_HEIGHT_CM = 2.0
SUBPROJECT_TABLE_TOP_SPACING_CM = 0.3
SUBPROJECT_CONTINUATION_TOP_SPACING_CM = 0.65
REMARK_FOOTER_GAP_CM = 1.0
REMARK_ROW_BASE_HEIGHT_CM = 0.65
REMARK_ROW_BASE_MAX_CM = 2.0
REMARK_ROW_LINE_HEIGHT_CM = 0.52
REMARK_ROW_MAX_NEWLINES = 16
REMARK_FILL_MIN_SPACE_CM = 2.5
REMARK_FILL_SAFETY_CM = 2.0
REMARK_BODY_HEIGHT_FACTOR = 1.25
REMARK_FILL_RATIO = 0.55
EMU_PER_CM = 360000
DEFAULT_TABLE_ROW_HEIGHT_CM = 0.48
LEGEND_IMAGE_MAX_WIDTH_CM = 16.2
LEGEND_IMAGE_MAX_HEIGHT_CM = 22.5
LEGEND_BROWSER_TIMEOUT_SECONDS = 20


def main():
    if len(sys.argv) < 4:
        raise SystemExit("Usage: python scripts/generate_formatted_report.py template.docx input.json output.docx")

    template_path = Path(sys.argv[1])
    input_path = Path(sys.argv[2])
    output_path = Path(sys.argv[3])
    data = json.loads(input_path.read_text(encoding="utf-8-sig"))

    buffer = build_formatted_report_docx(data, template_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_bytes(buffer.getvalue())


def ensure_blank_page_after_toc(document):
    body = document._element.body
    children = list(body)
    toc_index = None
    for index, child in enumerate(children):
        if child.tag != qn("w:sdt"):
            continue
        text = "".join(child.itertext())
        if "TOC" in text or "目录" in text.replace(" ", ""):
            toc_index = index
            break
    if toc_index is None:
        return

    next_index = toc_index + 1
    if next_index < len(children):
        next_child = children[next_index]
        if next_child.tag == qn("w:p"):
            for br in next_child.xpath(".//w:br"):
                if br.get(qn("w:type")) == "page":
                    return

    page_break_paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    break_element = OxmlElement("w:br")
    break_element.set(qn("w:type"), "page")
    run.append(break_element)
    page_break_paragraph.append(run)
    body.insert(next_index, page_break_paragraph)


def build_formatted_report_docx(data, template_path):
    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"报告模板不存在: {template_path}")

    document = Document(template_path)
    template_tables = [deepcopy(table._tbl) for table in document.tables]
    template_heading = first_report_heading_paragraph(document)

    update_cover_text(document, data)
    update_statement_text(document, data)
    fixed_tail = trim_trailing_blank_tail_elements(extract_fixed_tail_from_heading(document, "资质证书"))
    remove_report_body_after_first_chapter(document)
    ensure_blank_page_after_toc(document)
    toc_entries, next_page_no, body_section_index = append_report_body(document, data, template_tables, template_heading)
    blank_section_index = None
    if fixed_tail:
        certificate_title = chapter_heading(len(toc_entries) + 1, "资质证书")
        update_fixed_tail_heading(fixed_tail, "资质证书", certificate_title)
        cert_start_page = next_page_no
        if next_page_no % 2 == 0:
            blank_section_index = add_body_blank_page_section(document)
            next_page_no += 1
            cert_start_page = next_page_no
            add_body_continued_section(document)
            clear_section_footer(document.sections[blank_section_index])
        else:
            document.add_page_break()
        append_body_elements(document, fixed_tail)
        total_body_pages = cert_start_page + fixed_tail_page_count(fixed_tail) - 1
        if blank_section_index is not None:
            apply_body_section_footers_with_total(
                document,
                body_section_index,
                total_body_pages,
                skip_indices={blank_section_index},
            )
            clear_section_footer(document.sections[blank_section_index])
        toc_entries.append((certificate_title, cert_start_page))
    replace_toc_cache(document, toc_entries)
    if blank_section_index is not None:
        clear_section_footer(document.sections[blank_section_index])
    update_fields_on_open(document)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def update_cover_text(document, data):
    cover = cover_data(data)
    replacements = {
        "报告编号": cover_value(cover, "reportNumber", "reportNo", "reportCode"),
        "委托单位名称": cover_value(cover, "clientName", "entrustingUnit", "commissionUnit"),
        "受检项目名称": cover_value(cover, "projectName", "inspectedProjectName"),
        "受检项目地址": cover_value(cover, "projectAddress", "inspectedProjectAddress"),
        "联   系   人": cover_value(cover, "contactName", "contactPerson", "linkman"),
        "联 系 人": cover_value(cover, "contactName", "contactPerson", "linkman"),
        "联系人": cover_value(cover, "contactName", "contactPerson", "linkman"),
        "电        话": cover_value(cover, "phone", "contactPhone", "telephone"),
        "电    话": cover_value(cover, "phone", "contactPhone", "telephone"),
        "电话": cover_value(cover, "phone", "contactPhone", "telephone"),
        "本次检测时间": cover_value(cover, "currentInspectionDate", "inspectedDate", "inspectDate", "inspectionDate"),
        "下次检测时间": cover_value(cover, "nextInspectionDate", "nextDate", "nextInspectDate", "nextInspectionDate"),
        "检测机构名称": cover_value(cover, "inspectorName", "agencyName", "inspectionAgency", "detectAgency"),
        "检测机构地址": cover_value(cover, "inspectorAddress", "agencyAddress", "inspectionAgencyAddress"),
        "检测机构电话": cover_value(cover, "inspectorPhone", "agencyPhone", "inspectionAgencyPhone"),
    }
    cover_title = cover_value(cover, "coverTitle")
    cover_title_en = cover_value(cover, "coverTitleEn")
    report_book_info = cover_value(cover, "reportBookInfo")
    for paragraph in iter_all_paragraphs(document):
        text = paragraph.text.strip()
        compact_text = re.sub(r"\s+", "", text)
        if cover_title and compact_text in {"雷电防护装置检测报告", "防雷检测报告"}:
            set_paragraph_text_keep_format(paragraph, cover_title, paragraph.alignment)
            continue
        if cover_title_en and text.startswith("Inspection report"):
            set_paragraph_text_keep_format(paragraph, cover_title_en, paragraph.alignment)
            continue
        if report_book_info and "第" in text and "册" in text and "共" in text:
            set_paragraph_text_keep_format(paragraph, report_book_info, paragraph.alignment)
            continue
        if should_replace_cover_agency_title(compact_text, cover):
            set_paragraph_text_keep_format(paragraph, cover_value(cover, "inspectorName", "agencyName"), paragraph.alignment)
            continue
        for label, replacement in replacements.items():
            if text.startswith(label):
                matched_label = cover_matched_label(text, label)
                if label == "报告编号":
                    replace_cover_report_number(paragraph, matched_label, replacement or "")
                    continue
                replacement = cover_value_with_template_suffix(text, matched_label, replacement or "")
                replace_cover_paragraph_with_underlined_value(paragraph, matched_label, replacement)


def cover_data(data):
    cover = data.get("cover", {})
    return cover if isinstance(cover, dict) else {}


def should_replace_cover_agency_title(compact_text, cover):
    agency_name = cover_value(cover, "inspectorName", "agencyName")
    if not agency_name:
        return False
    if compact_text == re.sub(r"\s+", "", agency_name):
        return False
    if compact_text.startswith("检测机构"):
        return False
    return compact_text.endswith("公司") and ("检测" in compact_text or "防雷" in compact_text)


def update_statement_text(document, data):
    content = cover_value(cover_data(data), "statementContent")
    if not content:
        return

    paragraphs = document.paragraphs
    start_index = None
    for index, paragraph in enumerate(paragraphs):
        if re.sub(r"\s+", "", paragraph.text) == "声明":
            start_index = index + 1
            break
    if start_index is None:
        return

    end_index = len(paragraphs)
    for index in range(start_index, len(paragraphs)):
        text = paragraphs[index].text.strip()
        compact_text = re.sub(r"\s+", "", text)
        if compact_text in {"目录"} or compact_text.startswith("一、总表"):
            end_index = index
            break

    lines = statement_content_lines(content)
    slots = [paragraph for paragraph in paragraphs[start_index:end_index] if paragraph.text.strip()]
    for index, paragraph in enumerate(slots):
        text = lines[index] if index < len(lines) else ""
        remove_paragraph_numbering(paragraph)
        set_statement_paragraph_text_keep_format(paragraph, text, index)


def set_statement_paragraph_text_keep_format(paragraph, text, index):
    template_rpr = first_paragraph_run_properties(paragraph)
    alignment = paragraph.alignment
    paragraph_format = paragraph.paragraph_format
    space_before = paragraph_format.space_before
    space_after = paragraph_format.space_after
    clear_paragraph_runs(paragraph)
    paragraph.alignment = alignment
    statement_left_indent, statement_first_line_indent = statement_paragraph_indent(text, index)
    paragraph.paragraph_format.left_indent = statement_left_indent
    paragraph.paragraph_format.first_line_indent = statement_first_line_indent
    paragraph.paragraph_format.space_before = space_before or Pt(0)
    paragraph.paragraph_format.space_after = space_after or Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    if template_rpr is not None:
        run._r.insert(0, deepcopy(template_rpr))
    set_run_font(run, size=None, bold=False)


def statement_paragraph_indent(text, index):
    if re.match(r"^\d+[\.、]", text):
        return Pt(42.55), None
    if index >= 12:
        return Pt(24), None
    return None, None


def remove_paragraph_numbering(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def statement_content_lines(content):
    lines = []
    for raw_line in str(content).splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if re.sub(r"\s+", "", line) == "声明":
            continue
        line = re.sub(r"^\d+\s*(?=[一二三四五六七八九十]+、)", "", line)
        lines.append(line)
    return lines


def cover_matched_label(text, label):
    matched_label = label
    suffix = text[len(label) : len(label) + 1]
    if suffix in {"：", ":"}:
        matched_label += suffix
    return matched_label


def cover_value_with_template_suffix(text, label, value):
    existing_value = text[len(label) :].strip()
    for marker in ("（盖章）", "(盖章)"):
        if marker in existing_value and marker not in str(value):
            return f"{value}{marker}"
    return value


def cover_value(general, *keys):
    for key in keys:
        value = general.get(key)
        if value is not None and str(value).strip() != "":
            return value
    return ""


def replace_cover_report_number(paragraph, label, report_number):
    template_rpr = first_paragraph_run_properties(paragraph)
    clear_paragraph_runs(paragraph)

    label_run = paragraph.add_run(f"{label} ")
    if template_rpr is not None:
        label_run._r.insert(0, deepcopy(template_rpr))
    set_run_font(
        label_run,
        size=label_run.font.size.pt if label_run.font.size else None,
        bold=bool(label_run.bold) if label_run.bold is not None else None,
    )

    value_run = paragraph.add_run(str(report_number or ""))
    if template_rpr is not None:
        value_run._r.insert(0, deepcopy(template_rpr))
    set_run_font(
        value_run,
        size=value_run.font.size.pt if value_run.font.size else None,
        bold=bool(value_run.bold) if value_run.bold is not None else None,
    )
    value_run.font.color.rgb = RGBColor(255, 0, 0)


def iter_all_paragraphs(document):
    for paragraph_element in document._element.xpath(".//w:p"):
        yield paragraph_element_to_paragraph(paragraph_element)


def paragraph_element_to_paragraph(paragraph_element):
    from docx.text.paragraph import Paragraph

    return Paragraph(paragraph_element, None)


def replace_cover_paragraph_with_underlined_value(paragraph, label, replacement):
    template_rpr = first_paragraph_run_properties(paragraph)
    value, suffix = split_cover_suffix(str(replacement))
    for child in list(paragraph._p):
        if child.tag == qn("w:r"):
            paragraph._p.remove(child)
    set_cover_value_tab_stop(paragraph)

    label_run = paragraph.add_run(f"{label} ")
    if template_rpr is not None:
        label_run._r.insert(0, deepcopy(template_rpr))
    set_run_font(label_run, size=label_run.font.size.pt if label_run.font.size else None, bold=bool(label_run.bold) if label_run.bold is not None else None)

    value_run = paragraph.add_run(value)
    if template_rpr is not None:
        value_run._r.insert(0, deepcopy(template_rpr))
    set_run_font(value_run, size=value_run.font.size.pt if value_run.font.size else None, bold=bool(value_run.bold) if value_run.bold is not None else None)
    value_run.underline = True

    tab_run = paragraph.add_run()
    if template_rpr is not None:
        tab_run._r.insert(0, deepcopy(template_rpr))
    set_run_font(tab_run, size=tab_run.font.size.pt if tab_run.font.size else None, bold=bool(tab_run.bold) if tab_run.bold is not None else None)
    tab_run.underline = True
    tab_run._r.append(OxmlElement("w:tab"))

    if suffix:
        suffix_run = paragraph.add_run(suffix)
        if template_rpr is not None:
            suffix_run._r.insert(0, deepcopy(template_rpr))
        set_run_font(suffix_run, size=suffix_run.font.size.pt if suffix_run.font.size else None, bold=bool(suffix_run.bold) if suffix_run.bold is not None else None)


def split_cover_suffix(value):
    for marker in ("（盖章）", "(盖章)"):
        if marker in value:
            before, after = value.split(marker, 1)
            return before.rstrip(), marker + after
    return value.strip(), ""


def set_cover_value_tab_stop(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), COVER_UNDERLINE_TAB_POS)
    tabs.append(tab)


def update_fields_on_open(document):
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    for fld_char in document._element.xpath(".//w:fldChar[@w:fldCharType='begin']"):
        fld_char.set(qn("w:dirty"), "true")


def replace_toc_cache(document, toc_entries):
    for sdt in document._element.body.xpath("./w:sdt"):
        if "TOC " not in "".join(sdt.itertext()):
            continue
        content = sdt.find(qn("w:sdtContent"))
        if content is None:
            return
        for child in list(content):
            content.remove(child)
        content.append(build_toc_title_paragraph())
        for title, page_no in toc_entries:
            content.append(build_toc_entry_paragraph(title, page_no))
        return


def build_toc_title_paragraph():
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    p_pr.append(jc)
    paragraph.append(p_pr)
    run = toc_run("目    录", size=20, bold=True)
    paragraph.append(run)
    return paragraph


def build_toc_entry_paragraph(title, page_no):
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "8500")
    tabs.append(tab)
    p_pr.append(tabs)
    paragraph.append(p_pr)
    paragraph.append(toc_run(title, size=14))
    run_tab = OxmlElement("w:r")
    run_tab.append(OxmlElement("w:tab"))
    paragraph.append(run_tab)
    paragraph.append(toc_run(str(page_no), size=14))
    return paragraph


def toc_run(text, size=14, bold=False):
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_NAME)
    fonts.set(qn("w:hAnsi"), FONT_NAME)
    fonts.set(qn("w:eastAsia"), FONT_NAME)
    r_pr.append(fonts)
    if bold:
        r_pr.append(OxmlElement("w:b"))
        r_pr.append(OxmlElement("w:bCs"))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(size * 2))
    r_pr.append(sz_cs)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    return run


def remove_report_body_after_first_chapter(document):
    body = document._element.body
    children = list(body)
    start_index = None
    for index, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        text = "".join(child.itertext()).strip()
        if text.startswith("一、总表"):
            start_index = index
            break

    if start_index is None:
        return

    sect_pr = body.sectPr
    for child in children[start_index:]:
        if child.tag == qn("w:sectPr"):
            continue
        if child.getparent() is body:
            body.remove(child)
    if sect_pr is not None and sect_pr.getparent() is None:
        body.append(sect_pr)


def extract_fixed_tail_from_heading(document, heading_text):
    body = document._element.body
    children = list(body)
    start_index = None
    for index, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        text = "".join(child.itertext()).strip()
        if heading_text in text:
            start_index = index
            break
    if start_index is None:
        return []
    return [deepcopy(child) for child in children[start_index:] if child.tag != qn("w:sectPr")]


def trim_trailing_blank_tail_elements(elements):
    trimmed = list(elements)
    while trimmed and is_blank_tail_paragraph(trimmed[-1]):
        trimmed.pop()
    return trimmed


def is_blank_tail_paragraph(element):
    if element.tag != qn("w:p"):
        return False
    if "".join(element.itertext()).strip():
        return False
    if element.xpath(".//w:drawing") or element.xpath(".//w:pict") or element.xpath(".//w:object"):
        return False
    for br in element.xpath(".//w:br"):
        if br.get(qn("w:type")) == "page":
            return False
    return True


def append_body_elements(document, elements):
    body = document._element.body
    insert_index = len(body) - 1 if body.sectPr is not None else len(body)
    for element in elements:
        body.insert(insert_index, deepcopy(element))
        insert_index += 1


def update_fixed_tail_heading(elements, keyword, title):
    for element in elements:
        if element.tag != qn("w:p"):
            continue
        text = "".join(element.itertext()).strip()
        if keyword not in text:
            continue
        text_nodes = element.xpath(".//w:t")
        if not text_nodes:
            return
        text_nodes[0].text = title
        for node in text_nodes[1:]:
            node.text = ""
        return


def append_report_body(document, data, template_tables, template_heading):
    toc_entries = []
    page_no = 1
    chapter_index = 1

    body_section_index = len(document.sections) - 1
    start_report_body_section(document)
    heading = chapter_heading(chapter_index, "总表")
    add_heading(document, heading, template_heading, break_before=False)
    toc_entries.append((heading, page_no))
    summary_table = append_template_table(document, template_tables[0])
    fill_summary_table(summary_table, data)

    page_no += 1
    chapter_index += 1

    assistant = data.get("assistant", {})
    report_tables = data.get("reportTables", {})
    for chapter_key in report_body_chapter_order(data):
        if chapter_key == "overview":
            pages = add_subproject_section(
                document,
                chapter_heading(chapter_index, "子项目表（概况）"),
                assistant.get("overview", {}).get("projects", []),
                data,
                template_tables[1],
                template_heading,
                toc_entries,
                page_no,
            )
        elif chapter_key in {"power", "electronic"}:
            chapter_title = "子项目表（低压电源系统）" if chapter_key == "power" else "子项目表（电子信息系统）"
            pages = add_subproject_section(
                document,
                chapter_heading(chapter_index, chapter_title),
                assistant.get(chapter_key, {}).get("projects", []),
                data,
                template_tables[2],
                template_heading,
                toc_entries,
                page_no,
            )
        else:
            measurement_config = {
                "grounding": ("子项目表（接地电阻）", 3, "grounding"),
                "transition": ("子项目表（过渡电阻）", 4, "transition"),
                "spd": ("子项目表（SPD明细表）", 5, "spd"),
                "spdTest": ("子项目表（SPD测试表）", 6, "spd_test"),
            }.get(chapter_key)
            if measurement_config is None:
                continue
            chapter_title, template_index, measurement_kind = measurement_config
            if template_index >= len(template_tables):
                continue
            rows = report_tables.get(chapter_key, [])
            if not has_measurement_table_data(rows):
                continue
            pages = measurement_page_count(rows, template_tables[template_index], measurement_kind)
            heading = chapter_heading(chapter_index, chapter_title)
            add_heading(document, heading, template_heading)
            toc_entries.append((heading, page_no))
            add_template_measurement_tables(
                document,
                rows,
                template_tables[template_index],
                measurement_kind,
                data,
            )

        if pages:
            page_no += pages
            chapter_index += 1

    legend = data.get("legend", {})
    legend_page_count = legend_content_page_count(legend)
    if legend_page_count:
        heading = chapter_heading(chapter_index, "现场平面示意图")
        add_heading(document, heading, template_heading)
        toc_entries.append((heading, page_no))
        add_legend_tables(document, legend)
        page_no += legend_page_count
        chapter_index += 1

    return toc_entries, page_no, body_section_index


def report_body_chapter_order(data):
    legacy_order = ["overview", "power", "electronic", "grounding", "transition", "spd", "spdTest"]
    configured_order = data.get("legend", {}).get("reportChapterOrder")
    if not isinstance(configured_order, list) or not configured_order:
        return legacy_order

    supported = set(legacy_order)
    order = []
    for chapter_key in configured_order:
        if chapter_key in supported and chapter_key not in order:
            order.append(chapter_key)
    return order


def chapter_heading(index, title):
    numerals = "零一二三四五六七八九十"
    if index <= 10:
        prefix = numerals[index]
    elif index < 20:
        prefix = "十" + numerals[index - 10]
    else:
        tens, ones = divmod(index, 10)
        prefix = numerals[tens] + "十" + (numerals[ones] if ones else "")
    return f"{prefix}、{title}"


def data_with_summary_project_pages(data, template_tables):
    general = data.get("assistant", {}).get("general", {})
    project_items = general.get("projectItems", [])
    if not project_items:
        return data

    page_map = summary_project_page_map(data, template_tables)
    if not page_map:
        return data

    enriched = deepcopy(data)
    enriched_general = enriched.setdefault("assistant", {}).setdefault("general", {})
    enriched_items = []
    for item in project_items:
        enriched_item = deepcopy(item)
        item_name = str(item.get("name") or "").strip()
        if item_name in page_map:
            enriched_item["pages"] = page_map[item_name]
        enriched_items.append(enriched_item)
    enriched_general["projectItems"] = enriched_items
    return enriched


def summary_project_page_map(data, template_tables):
    page_map = {}
    page_no = 2

    for project in filtered_subprojects(data.get("assistant", {}).get("overview", {}).get("projects", [])):
        add_project_page(page_map, subproject_name(project, data), page_no)
        page_no += 1

    for project in filtered_subprojects(data.get("assistant", {}).get("power", {}).get("projects", [])):
        add_project_page(page_map, subproject_name(project, data), page_no)
        page_no += 1

    for project in filtered_subprojects(data.get("assistant", {}).get("electronic", {}).get("projects", [])):
        add_project_page(page_map, subproject_name(project, data), page_no)
        page_no += 1

    report_tables = data.get("reportTables", {})
    for kind, table_index in (("grounding", 3), ("transition", 4), ("spd", 5), ("spdTest", 6)):
        if table_index >= len(template_tables):
            continue
        rows = report_tables.get(kind, [])
        if not has_measurement_table_data(rows):
            continue
        groups = group_rows_by_place(rows)
        rows_per_table = measurement_rows_per_table(template_tables[table_index], "spd_test" if kind == "spdTest" else kind)
        for place_name, group_rows in groups:
            chunks = chunk_rows(group_rows, rows_per_table)
            if place_name:
                pages = [str(page_no + index) for index in range(len(chunks))]
                add_project_page(page_map, place_name, "、".join(pages))
            page_no += len(chunks)

    return page_map


def filtered_subprojects(projects):
    return [project for project in projects if project.get("rows")]


def add_project_page(page_map, name, page):
    name = str(name or "").strip()
    if not name:
        return
    pages = [item for item in str(page).split("、") if item]
    if name not in page_map:
        page_map[name] = "、".join(pages)
        return
    existing = [item for item in str(page_map[name]).split("、") if item]
    for page_item in pages:
        if page_item not in existing:
            existing.append(page_item)
    page_map[name] = "、".join(existing)


def append_template_table(document, table_element):
    tbl = deepcopy(table_element)
    document._element.body.insert(len(document._element.body) - 1, tbl)
    return document.tables[-1]


def first_report_heading_paragraph(document):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith("一、总表"):
            return deepcopy(paragraph._p)
    return None


def start_report_body_section(document):
    section = document.sections[-1]
    restart_section_page_numbering(section)
    write_body_section_footer(section)


def set_body_section_footer(section, total_pages=None):
    section.footer.is_linked_to_previous = False
    write_body_section_footer(section, total_pages=total_pages)


def write_body_section_footer(section, total_pages=None):
    paragraph = section.footer.paragraphs[0]
    clear_paragraph_runs(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_footer_text(paragraph, "第 ")
    add_field(paragraph, "PAGE")
    add_footer_text(paragraph, " 页 共 ")
    if total_pages is None:
        add_field(paragraph, "SECTIONPAGES")
    else:
        add_footer_text(paragraph, str(total_pages))
    add_footer_text(paragraph, " 页")


def strip_section_footer(section):
    sect_pr = section._sectPr
    for footer_ref in list(sect_pr.findall(qn("w:footerReference"))):
        sect_pr.remove(footer_ref)


def clear_section_footer(section):
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    clear_paragraph_runs(paragraph)


def continue_section_page_numbering(section):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is not None:
        sect_pr.remove(pg_num_type)


def add_body_blank_page_section(document):
    document.add_section(WD_SECTION.NEW_PAGE)
    section = document.sections[-1]
    continue_section_page_numbering(section)
    clear_section_footer(section)
    document.add_paragraph("")
    return len(document.sections) - 1


def add_body_continued_section(document):
    document.add_section(WD_SECTION.NEW_PAGE)
    section = document.sections[-1]
    continue_section_page_numbering(section)
    set_body_section_footer(section)


def apply_body_section_footers_with_total(document, body_section_index, total_pages, skip_indices=None):
    skip_indices = skip_indices or set()
    for index in range(body_section_index, len(document.sections)):
        if index in skip_indices:
            continue
        set_body_section_footer(document.sections[index], total_pages=total_pages)
    for index in skip_indices:
        clear_section_footer(document.sections[index])


def fixed_tail_page_count(fixed_tail):
    page_count = 1
    for element in fixed_tail:
        if element.tag != qn("w:p"):
            continue
        for br in element.xpath(".//w:br"):
            if br.get(qn("w:type")) == "page":
                page_count += 1
    return page_count


def restart_section_page_numbering(section):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), "1")


def clear_paragraph_runs(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_footer_text(paragraph, text):
    run = paragraph.add_run(text)
    set_run_font(run, size=9, bold=False)


def add_field(paragraph, field_code):
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)

    run_text = paragraph.add_run("1")
    set_run_font(run_text, size=9, bold=False)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def fill_summary_table(table, data):
    general = data.get("assistant", {}).get("general", {})
    project_items = meaningful_summary_items(
        general.get("projectItems", []),
        ("name", "lightningCategory", "lightningProtectionLevel", "pages"),
    )
    equipment = meaningful_summary_items(
        general.get("equipment", []),
        ("name", "model", "serial", "range", "calibrationDate"),
    )[:6]

    if len(table.rows) <= 21 and len(table.columns) >= 17:
        fill_summary_table_compact(table, general, project_items, equipment)
        trim_summary_optional_rows(table, len(equipment), len(project_items))
        return

    set_cell(table.cell(0, 3), summary_value(general.get("clientName")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(1, 3), summary_value(general.get("projectName")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(2, 3), summary_value(general.get("projectAddress")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(3, 3), summary_value(general.get("contactDepartment")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(3, 6), summary_value(general.get("contactName")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(3, 14), summary_value(general.get("contactPhone")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(4, 3), summary_value(general.get("inspectedDate")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(5, 3), summary_value(general.get("nextDate")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(5, 11), summary_value(general.get("detectionType")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(6, 3), summary_value(general.get("detectionBasis")), align=WD_ALIGN_PARAGRAPH.LEFT, empty_text="")

    for row_index in range(8, 14):
        item = equipment[row_index - 8] if row_index - 8 < len(equipment) else {}
        set_cell(table.cell(row_index, 1), summary_value(item.get("name")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 4), summary_value(item.get("model")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 7), summary_value(item.get("serial")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 10), summary_value(item.get("range")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 14), summary_value(item.get("calibrationDate")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")

    for row_index in range(16, 27):
        item = project_items[row_index - 16] if row_index - 16 < len(project_items) else {}
        set_cell(table.cell(row_index, 0), summary_value(item.get("id")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 2), summary_value(item.get("name")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 5), summary_value(item.get("lightningCategory")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 9), summary_value(item.get("lightningProtectionLevel")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 13), summary_value(item.get("pages")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")

    fill_summary_conclusion_cell(table.cell(27, 2), general.get("conclusion"))
    set_cell(table.cell(28, 2), signature_value(summary_value(general.get("tester"))), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(28, 8), signature_value(summary_value(general.get("reviewer"))), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(28, 15), signature_value(summary_value(general.get("signer"))), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    trim_summary_optional_rows(table, len(equipment), len(project_items))


def fill_summary_table_compact(table, general, project_items, equipment):
    set_cell(table.cell(0, 3), summary_value(general.get("clientName")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(1, 3), summary_value(general.get("projectName")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(2, 3), summary_value(general.get("projectAddress")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(3, 3), summary_value(general.get("contactDepartment")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(3, 8), summary_value(general.get("contactName")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(3, 15), summary_value(general.get("contactPhone")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(4, 3), summary_value(general.get("inspectedDate")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(5, 3), summary_value(general.get("nextDate")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(5, 12), summary_value(general.get("detectionType")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
    set_cell(table.cell(6, 3), summary_value(general.get("detectionBasis")), align=WD_ALIGN_PARAGRAPH.LEFT, empty_text="")

    for row_index in range(8, min(14, len(table.rows))):
        item = equipment[row_index - 8] if row_index - 8 < len(equipment) else {}
        set_cell(table.cell(row_index, 1), summary_value(item.get("name")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 4), summary_value(item.get("model")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 5), summary_value(item.get("serial")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 11), summary_value(item.get("range")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 15), summary_value(item.get("calibrationDate")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")

    for row_index in range(16, min(19, len(table.rows))):
        item = project_items[row_index - 16] if row_index - 16 < len(project_items) else {}
        set_cell(table.cell(row_index, 0), summary_value(item.get("id")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 2), summary_value(item.get("name")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 6), summary_value(item.get("lightningCategory")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 10), summary_value(item.get("lightningProtectionLevel")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, 14), summary_value(item.get("pages")), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")

    if len(table.rows) > 19:
        fill_summary_conclusion_cell(table.cell(19, 2), general.get("conclusion"))
    if len(table.rows) > 20:
        set_cell(table.cell(20, 2), signature_value(summary_value(general.get("tester"))), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(20, 9), signature_value(summary_value(general.get("reviewer"))), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(20, 16), signature_value(summary_value(general.get("signer"))), align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")


def meaningful_summary_items(items, content_keys):
    result = []
    for item in items or []:
        if not isinstance(item, dict):
            continue
        if any(has_value(item.get(key)) for key in content_keys):
            result.append(item)
    return result


def summary_value(value):
    return value if has_value(value) else None


def signature_value(value):
    if value is None:
        return None
    text = str(value).strip()
    compact_text = re.sub(r"\s+", "", text)
    if compact_text in {"检测人", "校核人", "授权签字人"}:
        return None
    if compact_text and set(compact_text) == {"?"}:
        return None
    return value


def fill_summary_conclusion_cell(cell, conclusion):
    paragraphs = ensure_cell_paragraph_count(cell, 11)
    intro, items = summary_conclusion_parts(conclusion)

    paragraph_specs = [
        ("", WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.LEFT),
        (intro, WD_ALIGN_PARAGRAPH.LEFT),
        (items[0] if len(items) > 0 else "", WD_ALIGN_PARAGRAPH.LEFT),
        (items[1] if len(items) > 1 else "", WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.LEFT),
        ("", WD_ALIGN_PARAGRAPH.RIGHT),
        ("", WD_ALIGN_PARAGRAPH.RIGHT),
        ("签发日期：       年     月     日", WD_ALIGN_PARAGRAPH.RIGHT),
    ]

    for paragraph, (text, alignment) in zip(paragraphs, paragraph_specs):
        set_paragraph_text_keep_format(paragraph, text, alignment)


def ensure_cell_paragraph_count(cell, count):
    while len(cell.paragraphs) < count:
        cell.add_paragraph()
    return cell.paragraphs[:count]


def set_paragraph_text_keep_format(paragraph, text, alignment):
    template_rpr = first_paragraph_run_properties(paragraph)
    clear_paragraph_runs(paragraph)
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text)
    if template_rpr is not None:
        run._r.insert(0, deepcopy(template_rpr))
    set_run_font(run, size=None, bold=False)


def summary_conclusion_parts(conclusion):
    if not has_value(conclusion):
        return "", ["", ""]

    if isinstance(conclusion, dict):
        intro = first_value(conclusion, "intro", "summary", "title", "prefix")
        items = conclusion.get("items") or conclusion.get("details") or conclusion.get("list") or []
        if isinstance(items, str):
            items = [line.strip() for line in items.splitlines() if line.strip()]
        if not isinstance(items, list):
            items = []
        items = [normalize_conclusion_item(item) for item in items if has_value(item)]
        return str(intro or ""), pad_conclusion_items(items)

    if isinstance(conclusion, list):
        items = [normalize_conclusion_item(item) for item in conclusion if has_value(item)]
        return "", pad_conclusion_items(items)

    text = str(conclusion).strip()
    text = re.sub(r"检测机构（检测专用章）.*", "", text)
    text = re.sub(r"签发日期：.*", "", text)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return "", ["", ""]

    intro = ""
    body_text = " ".join(lines)
    if "：" in body_text and body_text.startswith("经过"):
        intro, body_text = body_text.split("：", 1)
        intro = intro + "："
    elif len(lines) > 1 and not re.match(r"^[12][\.、\s]*", lines[0]):
        intro = lines[0]
        body_text = " ".join(lines[1:])

    numbered_parts = [item.strip() for item in re.split(r"(?<!\d)[12][\.、]\s*", body_text) if item.strip()]
    if len(numbered_parts) > 2:
        numbered_parts = numbered_parts[-2:]
    if len(numbered_parts) >= 2:
        items = [normalize_conclusion_item(item) for item in numbered_parts[:2]]
        return intro, pad_conclusion_items(items)

    sentences = [sentence.strip() for sentence in re.split(r"(?<=。)", body_text) if sentence.strip()]
    if not sentences:
        sentences = [body_text]
    items = [normalize_conclusion_item(sentence) for sentence in sentences[:2]]
    return intro, pad_conclusion_items(items)


def normalize_conclusion_item(text):
    value = re.sub(r"^[12][\.、]\s*", "", str(text).strip())
    return value


def first_value(source, *keys):
    for key in keys:
        value = source.get(key)
        if has_value(value):
            return value
    return ""


def pad_conclusion_items(items):
    padded = list(items)
    while len(padded) < 2:
        padded.append("")
    return padded[:2]


def trim_summary_optional_rows(table, equipment_count, project_count):
    if len(table.rows) <= 21:
        trim_compact_summary_optional_rows(table, equipment_count, project_count)
        return

    if project_count:
        for row_index in range(26, 16 + project_count - 1, -1):
            table._tbl.remove(table.rows[row_index]._tr)
    else:
        for row_index in range(26, 13, -1):
            table._tbl.remove(table.rows[row_index]._tr)

    if equipment_count:
        for row_index in range(13, 8 + equipment_count - 1, -1):
            table._tbl.remove(table.rows[row_index]._tr)
    else:
        for row_index in range(13, 6, -1):
            table._tbl.remove(table.rows[row_index]._tr)


def trim_compact_summary_optional_rows(table, equipment_count, project_count):
    if project_count:
        last_project_row = min(18, 16 + project_count - 1)
        for row_index in range(18, last_project_row, -1):
            if row_index < len(table.rows):
                table._tbl.remove(table.rows[row_index]._tr)
    else:
        for row_index in range(18, 13, -1):
            if row_index < len(table.rows):
                table._tbl.remove(table.rows[row_index]._tr)

    if equipment_count:
        last_equipment_row = min(13, 8 + equipment_count - 1)
        for row_index in range(13, last_equipment_row, -1):
            if row_index < len(table.rows):
                table._tbl.remove(table.rows[row_index]._tr)
    else:
        for row_index in range(13, 6, -1):
            if row_index < len(table.rows):
                table._tbl.remove(table.rows[row_index]._tr)


def add_subproject_section(document, heading, projects, data, template_table, template_heading, toc_entries=None, page_no=None):
    projects = [project for project in projects if project.get("rows")]
    if not projects:
        return 0

    add_heading(document, heading, template_heading)
    if toc_entries is not None and page_no is not None:
        toc_entries.append((heading, page_no))
    for index, project in enumerate(projects):
        if index:
            document.add_page_break()
            add_table_top_spacing(document)
        add_subproject_table(
            document,
            project,
            data,
            template_table,
            has_section_heading=index == 0,
        )
    return len(projects)


def format_detection_type(raw):
    if raw == "验收检测":
        return "☑验收检测  □定期检测"
    if raw == "定期检测":
        return "□验收检测  ☑定期检测"
    return raw


def add_subproject_table(document, project, data, template_table, has_section_heading=False):
    rows = project.get("rows", [])
    table = append_template_table(document, template_table)
    trim_subproject_template_rows(table, len(rows))
    expand_compact_subproject_category_column(table, rows)
    positions = subproject_positions(table)
    reset_subproject_template_merges(table, 2, len(rows), positions)
    row_meta = []

    set_cell(
        table.cell(0, positions["name"]),
        subproject_name(project, data),
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        empty_text="",
    )
    set_cell(
        table.cell(0, positions["date"]),
        project.get("inspectionDate"),
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        empty_text="",
    )

    for index, row_data in enumerate(rows):
        row_index = index + 2
        if row_index >= len(table.rows) - 1:
            break
        category = display_row_category(table, row_index, positions, row_data)
        subcategory = display_row_subcategory(table, row_index, positions, row_data)
        content = display_row_content(table, row_index, positions, row_data)
        standard = display_row_standard(row_data)

        row_meta.append((category, subcategory))
        if positions.get("category") is not None:
            category_cell = table.cell(row_index, positions["category"])
            set_cell(category_cell, category, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        if positions.get("subcategory") is not None:
            subcategory_cell = table.cell(row_index, positions["subcategory"])
            if positions.get("category") is None or subcategory_cell._tc is not table.cell(row_index, positions["category"])._tc:
                set_cell(subcategory_cell, subcategory, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        set_cell(table.cell(row_index, positions["content"]), content, bold=True, empty_text="")
        set_cell(table.cell(row_index, positions["standard"]), standard, empty_text="")
        set_cell(
            table.cell(row_index, positions["result"]),
            row_data.get("result"),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            empty_text="",
        )
        set_cell(
            table.cell(row_index, positions["conclusion"]),
            row_data.get("conclusion"),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            empty_text="",
        )
    merge_subproject_category_columns(table, 2, row_meta, positions)
    stretch_table_remark_row(
        document,
        table,
        has_section_heading=has_section_heading,
        is_continuation=not has_section_heading,
    )


def expand_compact_subproject_category_column(table, rows):
    """Restore the category/subcategory columns collapsed by the compact template."""
    if len(table.columns) != 7:
        return

    grid = table._tbl.tblGrid
    first_grid_column = grid.gridCol_lst[0]
    original_width = int(first_grid_column.get(qn("w:w")) or 855)
    category_width = max(1, original_width // 2)
    subcategory_width = max(1, original_width - category_width)
    first_grid_column.set(qn("w:w"), str(subcategory_width))
    category_grid_column = OxmlElement("w:gridCol")
    category_grid_column.set(qn("w:w"), str(category_width))
    grid.insert(0, category_grid_column)

    categories_with_subcategories = {
        str(row.get("category") or "").strip()
        for row in rows
        if str(row.get("subcategory") or "").strip()
        and str(row.get("subcategory") or "").strip() != str(row.get("category") or "").strip()
    }
    first_data_row = 2
    for row_index, row in enumerate(table.rows):
        first_cell = row._tr.tc_lst[0]
        if first_data_row <= row_index < first_data_row + len(rows):
            row_data = rows[row_index - first_data_row]
            category = str(row_data.get("category") or "").strip()
            if category in categories_with_subcategories:
                split_subproject_category_cell(
                    row,
                    first_cell,
                    category_width,
                    subcategory_width,
                )
                continue
        increase_cell_grid_span(first_cell, 1)


def split_subproject_category_cell(row, category_cell, category_width, subcategory_width):
    category_properties = category_cell.get_or_add_tcPr()
    set_xml_cell_width(category_properties, category_width)
    remove_xml_grid_span(category_properties)

    subcategory_cell = deepcopy(category_cell)
    subcategory_properties = subcategory_cell.get_or_add_tcPr()
    set_xml_cell_width(subcategory_properties, subcategory_width)
    remove_xml_grid_span(subcategory_properties)
    for vertical_merge in list(subcategory_properties.findall(qn("w:vMerge"))):
        subcategory_properties.remove(vertical_merge)
    for paragraph in list(subcategory_cell.findall(qn("w:p"))):
        subcategory_cell.remove(paragraph)
    subcategory_cell.append(OxmlElement("w:p"))
    category_cell.addnext(subcategory_cell)


def increase_cell_grid_span(cell_element, amount):
    properties = cell_element.get_or_add_tcPr()
    grid_span = properties.gridSpan
    current_span = int(grid_span.val) if grid_span is not None else 1
    if grid_span is None:
        grid_span = OxmlElement("w:gridSpan")
        properties.append(grid_span)
    grid_span.set(qn("w:val"), str(current_span + amount))


def remove_xml_grid_span(properties):
    for grid_span in list(properties.findall(qn("w:gridSpan"))):
        properties.remove(grid_span)


def set_xml_cell_width(properties, width):
    cell_width = properties.find(qn("w:tcW"))
    if cell_width is None:
        cell_width = OxmlElement("w:tcW")
        properties.insert(0, cell_width)
    cell_width.set(qn("w:w"), str(width))
    cell_width.set(qn("w:type"), "dxa")


def display_row_category(table, row_index, positions, row_data):
    value = row_data.get("category")
    if has_value(value):
        return value
    return value


def display_row_subcategory(table, row_index, positions, row_data):
    value = row_data.get("subcategory")
    if has_value(value):
        return value
    return value


def display_row_content(table, row_index, positions, row_data):
    content = row_data.get("content")
    if has_value(content):
        return content
    return content


def display_row_standard(row_data):
    standard = row_data.get("standard")
    if has_value(standard):
        return standard
    return standard


def template_cell_text(table, row_index, col_index):
    if col_index is None:
        return None
    try:
        return table.cell(row_index, col_index).text.strip()
    except IndexError:
        return None


def is_meaningful_template_label(text):
    return has_value(text) and not looks_like_standard_value(text)


def has_value(text):
    return text is not None and str(text).strip() not in {"", "-", "－", "–", "—"}


def has_measurement_table_data(rows):
    if not rows:
        return False
    return any(
        isinstance(row, dict) and any(has_value(value) for value in row.values())
        for row in rows
    )


def looks_like_standard_value(text):
    if not has_value(text):
        return False
    value_text = str(text).strip()
    standard_markers = (
        "GB",
        "JGJ",
        "≤",
        "≥",
        ">",
        "<",
        "宜",
        "应",
        "自然/人工",
        "独立接地",
        "共用接地",
        "埋地/架空",
        "TN-",
    )
    return any(marker in value_text for marker in standard_markers)


def subproject_positions(table):
    if len(table.columns) == 7:
        return {
            "name": 2,
            "date": 5,
            "category": 0,
            "subcategory": None,
            "content": 1,
            "standard": 3,
            "result": 4,
            "conclusion": 6,
        }
    return {
        "name": 3,
        "date": 6,
        "category": 0,
        "subcategory": 1,
        "content": 2,
        "standard": 4,
        "result": 5,
        "conclusion": 7,
    }


def trim_subproject_template_rows(table, data_row_count):
    first_data_row = 2
    remark_row_index = len(table.rows) - 1
    available_data_rows = remark_row_index - first_data_row
    if data_row_count > available_data_rows:
        for _ in range(data_row_count - available_data_rows):
            source_row = deepcopy(table.rows[remark_row_index - 1]._tr)
            table._tbl.insert(remark_row_index, source_row)
            remark_row_index += 1
    keep_until = first_data_row + data_row_count
    for row_index in range(remark_row_index - 1, keep_until - 1, -1):
        table._tbl.remove(table.rows[row_index]._tr)


def is_remark_table_row(row):
    if not row.cells:
        return False
    return row.cells[0].text.strip().startswith("备注")


def page_content_height_cm(document):
    section = document.sections[-1]
    return (section.page_height - section.top_margin - section.bottom_margin) / EMU_PER_CM


def page_content_top_offset_cm(has_section_heading=False, is_continuation=False):
    offset = SUBPROJECT_TABLE_TOP_SPACING_CM
    if has_section_heading:
        offset += SUBPROJECT_SECTION_HEADING_HEIGHT_CM
    if is_continuation:
        offset += SUBPROJECT_CONTINUATION_TOP_SPACING_CM
    return offset


def estimate_table_body_height_cm(table):
    total = 0.0
    for row_index, row in enumerate(table.rows):
        if row_index == len(table.rows) - 1 and is_remark_table_row(row):
            continue
        row_cm = row.height / EMU_PER_CM if row.height else 0.0
        content_lines = 1
        for cell in unique_row_cells(row):
            text = cell.text.strip()
            if text:
                content_lines = max(content_lines, estimate_remark_layout_lines(text))
        content_cm = 0.32 + (content_lines - 1) * 0.24
        if row_cm > 0:
            row_cm = max(row_cm, content_cm)
        else:
            row_cm = max(DEFAULT_TABLE_ROW_HEIGHT_CM, content_cm)
        total += row_cm
    return total * REMARK_BODY_HEIGHT_FACTOR


def estimate_remark_layout_lines(text, chars_per_line=20):
    lines = 0
    for part in str(text).splitlines() or [""]:
        length = len(part.strip())
        if length == 0:
            lines += 1
        else:
            lines += max(1, (length + chars_per_line - 1) // chars_per_line)
    return max(1, lines)


def calculate_remark_remaining_space_cm(document, table, has_section_heading=False, is_continuation=False):
    page_height = page_content_height_cm(document)
    top_offset = page_content_top_offset_cm(has_section_heading, is_continuation)
    body_height = estimate_table_body_height_cm(table)
    used = top_offset + body_height + REMARK_ROW_BASE_HEIGHT_CM + REMARK_FOOTER_GAP_CM + REMARK_FILL_SAFETY_CM
    return page_height - used


def remark_row_base_height_cm(remark_row):
    if remark_row.height:
        height_cm = remark_row.height / EMU_PER_CM
        if height_cm <= REMARK_ROW_BASE_MAX_CM:
            return height_cm
    return REMARK_ROW_BASE_HEIGHT_CM


def remark_row_newline_count(remaining_cm):
    if remaining_cm < REMARK_FILL_MIN_SPACE_CM:
        return 0
    fill_cm = remaining_cm * REMARK_FILL_RATIO
    return max(0, min(REMARK_ROW_MAX_NEWLINES, int(fill_cm / REMARK_ROW_LINE_HEIGHT_CM)))


def ensure_table_remark_row(table):
    if table.rows and is_remark_table_row(table.rows[-1]):
        return
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[-1])
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def stretch_table_remark_row(document, table, has_section_heading=False, is_continuation=False):
    ensure_table_remark_row(table)
    remark_row = table.rows[-1]
    base_height = remark_row_base_height_cm(remark_row)
    remaining = calculate_remark_remaining_space_cm(
        document,
        table,
        has_section_heading=has_section_heading,
        is_continuation=is_continuation,
    )
    newline_count = remark_row_newline_count(remaining)
    remark_row.height = Cm(base_height)
    remark_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for cell in unique_row_cells(remark_row):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell(
            cell,
            "备注：" + ("\n" * newline_count),
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            empty_text="",
            size=9,
            line_spacing=1,
        )


def merge_subproject_category_columns(table, first_data_row, row_meta, positions):
    if not row_meta:
        return

    category_col = positions.get("category")
    subcategory_col = positions.get("subcategory")
    if category_col is None:
        return

    index = 0
    while index < len(row_meta):
        category = row_meta[index][0]
        end = index
        while end + 1 < len(row_meta) and row_meta[end + 1][0] == category:
            end += 1

        row_start = first_data_row + index
        row_end = first_data_row + end
        if row_start != row_end:
            merged = safe_merge_vertical(table, row_start, row_end, category_col)
            set_cell(merged, category, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")
        else:
            set_cell(table.cell(row_start, category_col), category, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, empty_text="")

        if subcategory_col is not None:
            merge_subproject_subcategory_group(
                table,
                first_data_row,
                index,
                end,
                category_col,
                subcategory_col,
                row_meta,
            )
        index = end + 1


def merge_subproject_subcategory_group(
    table,
    first_data_row,
    start,
    end,
    category_col,
    subcategory_col,
    row_meta,
):
    sub_index = start
    while sub_index <= end:
        row_index = first_data_row + sub_index
        category_cell = table.cell(row_index, category_col)
        subcategory_cell = table.cell(row_index, subcategory_col)
        if subcategory_cell._tc is category_cell._tc:
            sub_index += 1
            continue

        subcategory = row_meta[sub_index][1]
        sub_end = sub_index
        while sub_end + 1 <= end and row_meta[sub_end + 1][1] == subcategory:
            next_row_index = first_data_row + sub_end + 1
            if table.cell(next_row_index, subcategory_col)._tc is table.cell(next_row_index, category_col)._tc:
                break
            sub_end += 1

        if sub_end > sub_index:
            merged = safe_merge_vertical(
                table,
                first_data_row + sub_index,
                first_data_row + sub_end,
                subcategory_col,
            )
        else:
            merged = subcategory_cell
        set_cell(
            merged,
            subcategory,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            empty_text="",
        )
        sub_index = sub_end + 1


def reset_subproject_template_merges(table, first_data_row, data_row_count, positions):
    if data_row_count <= 0:
        return

    columns = {
        positions.get("category"),
        positions.get("subcategory"),
    }
    columns.discard(None)
    last_data_row = min(first_data_row + data_row_count, len(table.rows) - 1)
    for row_index in range(first_data_row, last_data_row):
        row = table.rows[row_index]
        for column_index in columns:
            cell_element = physical_cell_at_grid_column(row, column_index)
            if cell_element is None:
                continue
            cell_properties = cell_element.get_or_add_tcPr()
            for vertical_merge in list(cell_properties.findall(qn("w:vMerge"))):
                cell_properties.remove(vertical_merge)


def physical_cell_at_grid_column(row, target_column):
    current_column = 0
    for cell_element in row._tr.tc_lst:
        cell_properties = cell_element.get_or_add_tcPr()
        grid_span = cell_properties.gridSpan
        span = int(grid_span.val) if grid_span is not None else 1
        if current_column <= target_column < current_column + span:
            return cell_element
        current_column += span
    return None


def safe_merge_vertical(table, row_start, row_end, col):
    try:
        return table.cell(row_start, col).merge(table.cell(row_end, col))
    except ValueError:
        return table.cell(row_start, col)


def add_remark_row(table, data_row_count):
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[-1])
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_sub_cell(cell, "备注：", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    total_rows_before_remark = data_row_count + 2
    if data_row_count >= 16:
        remark_height_cm = 0.35
    else:
        remark_height_cm = max(0.45, 13.8 - total_rows_before_remark * 0.42)
    row.height = Cm(remark_height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def subproject_name(project, data):
    return project.get("projectName")


def add_measurement_table(document, rows, columns):
    table = document.add_table(rows=1, cols=len(columns))
    table.autofit = True
    style_table(table)
    for idx, (_, header) in enumerate(columns):
        set_cell(table.cell(0, idx), header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, (key, _) in enumerate(columns):
            align = WD_ALIGN_PARAGRAPH.CENTER if key in {"marker", "standardValue", "measuredValue", "result"} else WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cells[idx], row_data.get(key), align=align)


def add_template_measurement_tables(document, rows, template_table, kind, data, include_empty=False):
    groups = group_rows_by_place(rows, include_empty=include_empty)
    rows_per_table = measurement_rows_per_table(template_table, kind)
    for table_index, (place_name, group_rows) in enumerate(groups):
        if table_index:
            document.add_page_break()
            add_measurement_continuation_spacing(document)
        page_capacity = measurement_group_capacity(len(group_rows), rows_per_table)
        add_template_measurement_table(
            document,
            group_rows,
            template_table,
            kind,
            data,
            place_name,
            page_capacity,
            has_section_heading=table_index == 0,
            is_continuation=table_index > 0,
        )


def measurement_page_count(rows, template_table, kind, include_empty=False):
    rows_per_table = measurement_rows_per_table(template_table, kind)
    page_count = 0
    for _, group_rows in group_rows_by_place(rows, include_empty=include_empty):
        page_count += max(1, (len(group_rows) + rows_per_table - 1) // rows_per_table)
    return page_count


def measurement_group_capacity(data_row_count, rows_per_page):
    data_row_count = max(0, data_row_count)
    if data_row_count <= rows_per_page:
        return rows_per_page
    return data_row_count


def add_measurement_continuation_spacing(document):
    add_table_top_spacing(document)


def add_table_top_spacing(document):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(18)


def add_template_measurement_table(
    document,
    group_rows,
    template_table,
    kind,
    data,
    place_name,
    table_capacity,
    has_section_heading=False,
    is_continuation=False,
):
    table = append_template_table(document, template_table)
    trim_measurement_rows_to_fit(table, len(group_rows), table_capacity, kind)
    prepare_measurement_template_rows(table, kind)
    repeat_measurement_header_rows(table, kind)
    format_measurement_conductor_spec_column(table, kind)
    positions = measurement_positions(kind)
    date_value = measurement_date(data) if group_rows else None
    set_cell(table.cell(0, positions["name"]), place_name, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.cell(0, positions["date"]), date_value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    if kind == "spd_test":
        first_data_row = measurement_first_data_row(kind)
        for item_index, row_data in enumerate(group_rows):
            row_index = first_data_row + item_index * 2
            if row_index >= len(table.rows) - 1:
                break
            fill_spd_test_row_pair(table, row_index, row_data)
    else:
        for row_index, row_data in enumerate(group_rows, start=measurement_first_data_row(kind)):
            if row_index >= len(table.rows) - 1:
                break
            fill_measurement_row(table, row_index, row_data, kind)
    stretch_table_remark_row(
        document,
        table,
        has_section_heading=has_section_heading,
        is_continuation=is_continuation,
    )


def chunk_rows(rows, size):
    if not rows:
        return [[]]
    return [rows[index : index + size] for index in range(0, len(rows), size)]


def measurement_rows_per_table(template_table, kind):
    if kind == "spd_test":
        return 15
    if kind in {"grounding", "transition", "spd"}:
        return 30
    template_row_count = len(template_table.xpath("./w:tr"))
    return max(1, template_row_count - 4)


def trim_measurement_rows_to_fit(table, data_row_count, max_data_rows, kind=None):
    first_data_row = measurement_first_data_row(kind)
    row_multiplier = 2 if kind == "spd_test" else 1
    keep_data_rows = max(max(0, data_row_count), max_data_rows) * row_multiplier
    ensure_measurement_data_row_capacity(table, first_data_row, keep_data_rows)
    remark_row_index = len(table.rows) - 1
    if table.rows and is_remark_table_row(table.rows[-1]):
        remark_row_index -= 1
    keep_until = min(first_data_row + keep_data_rows, remark_row_index)
    for row_index in range(remark_row_index - 1, keep_until - 1, -1):
        table._tbl.remove(table.rows[row_index]._tr)


def repeat_measurement_header_rows(table, kind=None):
    for row in table.rows[:measurement_first_data_row(kind)]:
        properties = row._tr.get_or_add_trPr()
        header = properties.find(qn("w:tblHeader"))
        if header is None:
            header = OxmlElement("w:tblHeader")
            properties.append(header)
        header.set(qn("w:val"), "true")


def ensure_measurement_data_row_capacity(table, first_data_row, required_data_rows):
    required_total_rows = first_data_row + required_data_rows + 1
    while len(table.rows) < required_total_rows:
        insert_index = len(table.rows)
        if table.rows and is_remark_table_row(table.rows[-1]):
            insert_index = len(table.rows) - 1
        source_row_index = max(first_data_row, insert_index - 1)
        cloned_row = deepcopy(table.rows[source_row_index]._tr)
        table._tbl.insert(insert_index, cloned_row)


def group_rows_by_place(rows, include_empty=False):
    if not rows and include_empty:
        return [("", [])]
    groups = []
    indexes = {}
    for row in rows:
        place_name = row.get("placeName") or ""
        if place_name not in indexes:
            indexes[place_name] = len(groups)
            groups.append((place_name, []))
        groups[indexes[place_name]][1].append(row)
    return groups


def measurement_date(data):
    raw = (
        data.get("legend", {}).get("inspectionDateValue")
        or data.get("legend", {}).get("inspectionDate")
        or data.get("assistant", {}).get("general", {}).get("inspectedDate")
    )
    return format_chinese_date(raw)


def format_chinese_date(raw):
    if not raw:
        return raw
    text = str(raw).strip()
    for fmt in ("%Y-%m-%d", "%Y/%m/%d"):
        try:
            date = datetime.strptime(text, fmt)
            return f"{date.year}年{date.month}月{date.day}日"
        except ValueError:
            pass
    return text


def measurement_positions(kind):
    if kind == "grounding":
        return {"name": 2, "date": 6}
    if kind == "transition":
        return {"name": 2, "date": 6}
    if kind == "spd":
        return {"name": 2, "date": 6}
    if kind == "spd_test":
        return {"name": 2, "date": 7}
    return {"name": 2, "date": 6}


def fill_measurement_row(table, row_index, row_data, kind):
    if kind == "grounding":
        set_cell(table.cell(row_index, 0), row_data.get("marker"), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(table.cell(row_index, 1), join_nonempty(row_data.get("workLocation"), row_data.get("equipmentName")))
        set_cell(table.cell(row_index, 3), row_data.get("conductorSpec"), align=WD_ALIGN_PARAGRAPH.CENTER, size=6.5)
        set_cell(table.cell(row_index, 4), row_data.get("protectionZone"), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(table.cell(row_index, 5), row_data.get("standardValue"), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(table.cell(row_index, 6), row_data.get("measuredValue"), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(table.cell(row_index, 7), row_data.get("result"), align=WD_ALIGN_PARAGRAPH.CENTER)
    elif kind == "transition":
        set_cell(table.cell(row_index, 0), row_data.get("marker"), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(
            table.cell(row_index, 1),
            join_nonempty(row_data.get("workLocation"), row_data.get("equipmentName"), row_data.get("referencePoint")),
        )
        set_cell(table.cell(row_index, 3), row_data.get("conductorSpec"), align=WD_ALIGN_PARAGRAPH.CENTER, size=6.5)
        set_cell(table.cell(row_index, 4), row_data.get("protectionZone"), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(table.cell(row_index, 5), row_data.get("standardValue"), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(table.cell(row_index, 6), row_data.get("measuredValue"), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(table.cell(row_index, 7), row_data.get("result"), align=WD_ALIGN_PARAGRAPH.CENTER)
    elif kind == "spd":
        set_cell(table.cell(row_index, 0), row_data.get("marker"), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(table.cell(row_index, 1), row_data.get("spdModel"))
        set_cell(table.cell(row_index, 2), row_data.get("installLocation"))
        set_cell(table.cell(row_index, 3), row_data.get("wireLength"), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(table.cell(row_index, 4), row_data.get("spdLevel"), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(table.cell(row_index, 5), row_data.get("installQuantity"), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(table.cell(row_index, 6), row_data.get("measuredValue"), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(table.cell(row_index, 7), row_data.get("result"), align=WD_ALIGN_PARAGRAPH.CENTER)
    elif kind == "spd_test":
        fill_spd_test_row_pair(table, row_index, row_data)


def fill_spd_test_row_pair(table, row_index, row_data):
    set_cell(table.cell(row_index, 0), first_row_value(row_data, "marker", "id", "code"), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.cell(row_index, 1), first_row_value(row_data, "spdModel", "model"))
    set_cell(table.cell(row_index, 2), first_row_value(row_data, "installLocation", "workLocation", "location"))
    set_cell(table.cell(row_index, 3), first_row_value(row_data, "u1maL1", "voltageL1", "ucL1"), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.cell(row_index, 4), first_row_value(row_data, "u1maL2", "voltageL2", "ucL2"), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.cell(row_index, 5), first_row_value(row_data, "leakageL1", "leakageCurrentL1"), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.cell(row_index, 6), first_row_value(row_data, "leakageL2", "leakageCurrentL2"), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.cell(row_index, 7), first_row_value(row_data, "insulationL1", "insulationResistanceL1"), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.cell(row_index, 8), first_row_value(row_data, "insulationL2", "insulationResistanceL2"), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.cell(row_index, 9), row_data.get("result"), align=WD_ALIGN_PARAGRAPH.CENTER)

    next_row_index = row_index + 1
    if next_row_index >= len(table.rows) - 1:
        return
    set_cell(table.cell(next_row_index, 3), first_row_value(row_data, "u1maL3", "voltageL3", "ucL3"), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.cell(next_row_index, 4), first_row_value(row_data, "u1maN", "voltageN", "ucN"), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.cell(next_row_index, 5), first_row_value(row_data, "leakageL3", "leakageCurrentL3"), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.cell(next_row_index, 6), first_row_value(row_data, "leakageN", "leakageCurrentN"), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.cell(next_row_index, 7), first_row_value(row_data, "insulationL3", "insulationResistanceL3"), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.cell(next_row_index, 8), first_row_value(row_data, "insulationN", "insulationResistanceN"), align=WD_ALIGN_PARAGRAPH.CENTER)


def first_row_value(row_data, *keys):
    for key in keys:
        value = row_data.get(key)
        if has_value(value):
            return value
    return None


def trim_measurement_template_rows(table, data_row_count):
    first_data_row = 2
    remark_row_index = len(table.rows) - 1
    keep_until = first_data_row + data_row_count
    for row_index in range(remark_row_index - 1, keep_until - 1, -1):
        table._tbl.remove(table.rows[row_index]._tr)


def prepare_measurement_template_rows(table, kind=None):
    last_index = len(table.rows) - 1
    if table.rows and is_remark_table_row(table.rows[-1]):
        last_index -= 1
    for row_index in range(measurement_first_data_row(kind), last_index):
        clear_measurement_data_row(table, row_index)


def format_measurement_conductor_spec_column(table, kind):
    if kind not in {"grounding", "transition"}:
        return
    column_index = 3
    for row_index, row in enumerate(table.rows):
        if column_index >= len(row.cells):
            continue
        cell = row.cells[column_index]
        set_cell_margins(cell, top=10, start=15, bottom=10, end=15)
        if row_index < measurement_first_data_row(kind):
            text = re.sub(r"\s+", "", cell.text)
            if "连接导体材料规格" in text:
                set_cell(cell, "连接导体材料规格", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=8)


def measurement_first_data_row(kind):
    return 4 if kind == "spd_test" else 2


def clear_measurement_data_row(table, row_index):
    for cell in unique_row_cells(table.rows[row_index]):
        set_cell(cell, None, align=WD_ALIGN_PARAGRAPH.CENTER)


def unique_row_cells(row):
    seen = set()
    cells = []
    for cell in row.cells:
        cell_id = id(cell._tc)
        if cell_id in seen:
            continue
        seen.add(cell_id)
        cells.append(cell)
    return cells


def join_nonempty(*items):
    return "，".join(str(item) for item in items if item is not None and str(item) != "")


def legend_content_page_count(legend):
    if not legend:
        return 0
    images = legend_image_sources(legend)
    if images:
        return len(images)
    canvases = extract_legend_canvases(legend)
    if canvases:
        return len(canvases)
    return 1 if legend.get("testPoints") else 0


def add_legend_tables(document, legend):
    rendered_images = render_legend_images(legend)
    if rendered_images:
        for index, image in enumerate(rendered_images):
            if index:
                document.add_page_break()
            add_legend_image(document, image)
        return

    test_points = legend.get("testPoints", [])
    if test_points:
        add_subheading(document, "检测点位")
        rows = []
        for item in test_points:
            fields = item.get("reportFields", {})
            rows.append(
                {
                    "label": item.get("label"),
                    "reportType": item.get("reportType"),
                    "equipmentName": fields.get("equipmentName"),
                    "workLocation": fields.get("workLocation"),
                    "standardValue": fields.get("standardValue"),
                    "measuredValue": fields.get("measuredValue"),
                    "result": fields.get("result"),
                }
            )
        add_measurement_table(
            document,
            rows,
            [
                ("label", "点位编号"),
                ("reportType", "检测类型"),
                ("equipmentName", "设备名称"),
                ("workLocation", "位置"),
                ("standardValue", "标准值"),
                ("measuredValue", "实测值"),
                ("result", "结论"),
            ],
        )


def render_legend_images(legend):
    images = []
    for source in legend_image_sources(legend):
        image = load_legend_image_source(source)
        if image:
            images.append(image)
    if images:
        return images

    for canvas in extract_legend_canvases(legend):
        svg_text = build_legend_svg(canvas)
        if not svg_text:
            continue
        width, height = legend_canvas_dimensions(canvas)
        image = svg_text_to_png_image(svg_text, width, height)
        if image:
            images.append(image)
            continue
        images.append(
            {
                "stream": BytesIO(svg_text.encode("utf-8")),
                "format": "svg",
                "width": width,
                "height": height,
            }
        )
    return images


def legend_image_sources(legend):
    sources = []
    for key in (
        "image",
        "imageData",
        "imageDataUrl",
        "imageBase64",
        "png",
        "pngData",
        "pngDataUrl",
        "screenshot",
        "previewImage",
    ):
        source = legend.get(key)
        if source:
            sources.append(source)
    return sources


def load_legend_image_source(source):
    if not isinstance(source, str):
        return None
    text = source.strip()
    if not text:
        return None

    if text.startswith("data:"):
        header, payload = text.split(",", 1)
        mime_type = header.split(";", 1)[0][5:].lower()
        if ";base64" in header:
            raw = b64decode(payload)
        else:
            raw = unquote(payload).encode("utf-8")
        image_format = "svg" if mime_type == "image/svg+xml" else "raster"
        return {"stream": BytesIO(raw), "format": image_format, "width": 900, "height": 600}

    if text.lstrip().startswith("<svg"):
        return {"stream": BytesIO(text.encode("utf-8")), "format": "svg", "width": 900, "height": 600}

    possible_path = Path(text)
    if possible_path.exists() and possible_path.is_file():
        raw = possible_path.read_bytes()
        image_format = "svg" if possible_path.suffix.lower() == ".svg" else "raster"
        return {"stream": BytesIO(raw), "format": image_format, "width": 900, "height": 600}

    try:
        raw = b64decode(text)
    except Exception:
        return None
    return {"stream": BytesIO(raw), "format": "raster", "width": 900, "height": 600}


def extract_legend_canvases(legend):
    workspace = normalize_legend_workspace(legend)
    if not workspace:
        return []

    if workspace.get("boardWidth") and workspace.get("boardHeight"):
        return [workspace]

    if has_canvas_content(workspace):
        return [workspace]

    tab_data = workspace.get("tabData") or {}
    if not isinstance(tab_data, dict) or not tab_data:
        return []

    if workspace.get("exportAllTabs"):
        return [canvas for canvas in tab_data.values() if isinstance(canvas, dict) and has_canvas_content(canvas)]

    active_tab_id = workspace.get("activeTabId")
    if active_tab_id is not None:
        canvas = tab_data.get(str(active_tab_id))
        if isinstance(canvas, dict) and has_canvas_content(canvas):
            return [canvas]

    for canvas in tab_data.values():
        if isinstance(canvas, dict) and has_canvas_content(canvas):
            return [canvas]
    return []


def has_canvas_content(canvas):
    return bool(
        isinstance(canvas, dict)
        and (
            canvas.get("importedBoard")
            or canvas.get("blocks")
            or canvas.get("paths")
            or canvas.get("texts")
            or canvas.get("testPoints")
        )
    )


def normalize_legend_workspace(legend):
    if not isinstance(legend, dict):
        return None
    for key in ("workspace", "drawingWorkspace", "drawingData", "legendWorkspace"):
        nested = legend.get(key)
        if isinstance(nested, dict):
            return nested
    return legend


def build_legend_svg(canvas):
    width, height = legend_canvas_dimensions(canvas)
    view_box = legend_canvas_view_box(canvas)
    if width <= 0 or height <= 0:
        return ""

    parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width:g}" height="{height:g}" viewBox="{view_box}">',
        "<style>",
        "text{font-family:'Microsoft YaHei','SimSun',Arial,sans-serif;fill:#111}",
        ".legend-path{fill:none;stroke:#222;stroke-width:1.2}",
        ".legend-point{fill:#000;stroke:#000;stroke-width:1}",
        "</style>",
        '<rect x="0" y="0" width="100%" height="100%" fill="#fff"/>',
    ]

    imported_board = canvas.get("importedBoard")
    if isinstance(imported_board, dict):
        parts.append(nested_svg_element(imported_board, 0, 0, width, height))

    for path in canvas.get("paths") or []:
        parts.append(path_to_svg(path))

    for index, block in enumerate(canvas.get("blocks") or []):
        parts.append(nested_svg_element(block, block.get("x", 0), block.get("y", 0), block.get("width", 40), block.get("height", 40), block.get("rotation", 0), f"block{index}_"))

    for text in canvas.get("texts") or []:
        parts.append(text_to_svg(text))

    for point in canvas.get("testPoints") or []:
        parts.append(test_point_to_svg(point))

    parts.append("</svg>")
    return "".join(part for part in parts if part)


def legend_canvas_dimensions(canvas):
    if canvas.get("boardWidth") and canvas.get("boardHeight"):
        return float(canvas.get("boardWidth")), float(canvas.get("boardHeight"))
    if canvas.get("width") and canvas.get("height"):
        return float(canvas.get("width")), float(canvas.get("height"))
    min_x, min_y, max_x, max_y = legend_canvas_bounds(canvas)
    return max(max_x - min_x + 40, 1), max(max_y - min_y + 40, 1)


def legend_canvas_view_box(canvas):
    if canvas.get("boardWidth") and canvas.get("boardHeight"):
        return f'0 0 {float(canvas.get("boardWidth")):g} {float(canvas.get("boardHeight")):g}'
    if canvas.get("width") and canvas.get("height"):
        return f'0 0 {float(canvas.get("width")):g} {float(canvas.get("height")):g}'
    min_x, min_y, max_x, max_y = legend_canvas_bounds(canvas)
    return f"{min_x - 20:g} {min_y - 20:g} {max(max_x - min_x + 40, 1):g} {max(max_y - min_y + 40, 1):g}"


def legend_canvas_bounds(canvas):
    xs = []
    ys = []

    def add_rect(x, y, width=0, height=0):
        x = float(x or 0)
        y = float(y or 0)
        width = float(width or 0)
        height = float(height or 0)
        xs.extend([x, x + width])
        ys.extend([y, y + height])

    imported_board = canvas.get("importedBoard")
    if isinstance(imported_board, dict):
        add_rect(0, 0, imported_board.get("width") or canvas.get("boardWidth"), imported_board.get("height") or canvas.get("boardHeight"))

    for block in canvas.get("blocks") or []:
        add_rect(block.get("x"), block.get("y"), block.get("width"), block.get("height"))

    for text in canvas.get("texts") or []:
        add_rect(text.get("x"), text.get("y"), text.get("width"), text.get("height"))

    for point in canvas.get("testPoints") or []:
        x = float(point.get("x") or 0)
        y = float(point.get("y") or 0)
        xs.extend([x - 12, x + 35])
        ys.extend([y - 18, y + 22])

    for path in canvas.get("paths") or []:
        for point in path.get("points") or []:
            xs.append(float(point.get("x") or 0))
            ys.append(float(point.get("y") or 0))

    if not xs or not ys:
        return 0, 0, 900, 600
    return min(xs), min(ys), max(xs), max(ys)


def nested_svg_element(item, x, y, width, height, rotation=0, prefix="nested_"):
    raw_svg = item.get("svg")
    if not raw_svg:
        return ""
    x = float(x or 0)
    y = float(y or 0)
    width = float(width or item.get("width") or 40)
    height = float(height or item.get("height") or 40)
    rotation = float(rotation or 0)
    transform = f"translate({x:g} {y:g})"
    if rotation:
        transform += f" rotate({rotation:g} {width / 2:g} {height / 2:g})"
    return f'<g transform="{transform}">{place_svg_at_size(namespace_svg_ids(raw_svg, prefix), width, height)}</g>'


def namespace_svg_ids(raw_svg, prefix):
    svg = raw_svg
    ids = re.findall(r'\bid="([^"]+)"', svg)
    for element_id in ids:
        safe_id = prefix + re.sub(r"[^A-Za-z0-9_.:-]", "_", element_id)
        svg = svg.replace(f'id="{element_id}"', f'id="{safe_id}"')
        svg = svg.replace(f'url(#{element_id})', f'url(#{safe_id})')
        svg = svg.replace(f'href="#{element_id}"', f'href="#{safe_id}"')
        svg = svg.replace(f'xlink:href="#{element_id}"', f'xlink:href="#{safe_id}"')
    return svg


def place_svg_at_size(raw_svg, width, height):
    svg = raw_svg.strip()
    svg = re.sub(r"^<\?xml[^>]*>\s*", "", svg)
    match = re.match(r"(<svg\b)([^>]*)(>)", svg)
    if not match:
        return svg
    attrs = re.sub(r'\s(?:x|y|width|height)="[^"]*"', "", match.group(2))
    opening = f'{match.group(1)} x="0" y="0" width="{width:g}" height="{height:g}"{attrs}{match.group(3)}'
    return opening + svg[match.end() :]


def path_to_svg(path):
    points = path.get("points") or []
    if not points:
        return ""
    point_text = " ".join(f'{float(point.get("x", 0)):g},{float(point.get("y", 0)):g}' for point in points)
    if path.get("closed"):
        return f'<polygon class="legend-path" points="{point_text}"/>'
    return f'<polyline class="legend-path" points="{point_text}"/>'


def text_to_svg(text):
    content = escape(str(text.get("text") or text.get("name") or ""))
    if not content:
        return ""
    x = float(text.get("x") or 0)
    y = float(text.get("y") or 0)
    size = float(text.get("fontSize") or 14)
    width = float(text.get("width") or 0)
    height = float(text.get("height") or 0)
    orientation = text.get("orientation")
    attrs = f'x="{x:g}" y="{y + size:g}" font-size="{size:g}"'
    if orientation == "vertical":
        cx = x + width / 2
        cy = y + height / 2
        attrs += f' transform="rotate(90 {cx:g} {cy:g})"'
    return f"<text {attrs}>{content}</text>"


def test_point_to_svg(point):
    x = float(point.get("x") or 0)
    y = float(point.get("y") or 0)
    size = max(float(point.get("size") or 1), 1) * 5
    label = escape(str(point.get("label") or ""))
    side = point.get("side") or "right"
    rotations = {"top": 0, "right": 90, "bottom": 180, "left": 270}
    rotation = rotations.get(side, 90)
    label_dx = 7 if side != "left" else -22
    label_dy = -7 if side in {"top", "right"} else 14
    return (
        f'<g transform="translate({x:g} {y:g}) rotate({rotation:g})">'
        f'<polygon class="legend-point" points="0,{-size:g} {size:g},{size:g} {-size:g},{size:g}"/>'
        "</g>"
        f'<text x="{x + label_dx:g}" y="{y + label_dy:g}" font-size="8">{label}</text>'
    )


def svg_text_to_png_image(svg_text, width, height):
    browser_path = find_headless_browser()
    if not browser_path:
        return None

    window_width = max(int(float(width)), 1)
    window_height = max(int(float(height)), 1)
    with TemporaryDirectory(prefix="legend-render-") as temp_dir:
        temp_path = Path(temp_dir)
        svg_path = temp_path / "legend.svg"
        png_path = temp_path / "legend.png"
        svg_path.write_text(svg_text, encoding="utf-8")
        command = [
            str(browser_path),
            "--headless",
            "--disable-gpu",
            "--hide-scrollbars",
            "--disable-dev-shm-usage",
            f"--window-size={window_width},{window_height}",
            f"--screenshot={png_path}",
            svg_path.as_uri(),
        ]
        result = subprocess.run(
            command,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=LEGEND_BROWSER_TIMEOUT_SECONDS,
            check=False,
        )
        if result.returncode != 0 or not png_path.exists():
            return None
        png_bytes = png_path.read_bytes()

    return {
        "stream": BytesIO(png_bytes),
        "format": "raster",
        "width": width,
        "height": height,
    }


def find_headless_browser():
    candidates = [
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return path
    return None


def add_legend_image(document, image):
    image = rotate_legend_image_for_page(image)
    image_stream = image["stream"]
    width = image.get("width") or 900
    height = image.get("height") or 600
    image_stream.seek(0)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(0)

    display_width = LEGEND_IMAGE_MAX_WIDTH_CM
    display_height = display_width * float(height or 1) / float(width or 1)
    if display_height > LEGEND_IMAGE_MAX_HEIGHT_CM:
        display_height = LEGEND_IMAGE_MAX_HEIGHT_CM
        display_width = display_height * float(width or 1) / float(height or 1)

    run = paragraph.add_run()
    if image.get("format") == "svg":
        add_svg_picture(document, run, image_stream.getvalue(), display_width, display_height)
    else:
        run.add_picture(image_stream, width=Cm(display_width), height=Cm(display_height))


def rotate_legend_image_for_page(image):
    if image.get("format") == "svg":
        return image
    try:
        from PIL import Image
    except ImportError:
        return image

    source = image["stream"]
    source.seek(0)
    try:
        with Image.open(source) as raw:
            rotated = raw.convert("RGBA").rotate(90, expand=True)
            output = BytesIO()
            rotated.save(output, format="PNG")
    except Exception:
        source.seek(0)
        return image

    output.seek(0)
    rotated_image = dict(image)
    rotated_image["stream"] = output
    rotated_image["format"] = "raster"
    rotated_image["width"] = image.get("height") or rotated.width
    rotated_image["height"] = image.get("width") or rotated.height
    return rotated_image


def add_svg_picture(document, run, svg_bytes, width_cm, height_cm):
    package = document.part.package
    partname = package.next_partname("/word/media/legend%d.svg")
    image_part = Part(partname, "image/svg+xml", svg_bytes, package)
    relationship_id = document.part.relate_to(image_part, RT.IMAGE)
    shape_id = document.part.next_id
    inline = CT_Inline.new_pic_inline(shape_id, relationship_id, partname.filename, Cm(width_cm), Cm(height_cm))
    run._r.add_drawing(inline)


def add_heading(document, text, template_heading=None, break_before=True):
    if break_before:
        document.add_page_break()
    if template_heading is not None:
        paragraph_element = deepcopy(template_heading)
        document._element.body.insert(len(document._element.body) - 1, paragraph_element)
        paragraph = document.paragraphs[-1]
        replace_paragraph_text_keep_style(paragraph, text)
    else:
        paragraph = document.add_paragraph(style="Heading 1")
        replace_paragraph_text_keep_style(paragraph, text)


def add_subheading(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=True)


def add_spacing(document, points):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(points)


def replace_paragraph_text_keep_style(paragraph, text):
    template_rpr = first_paragraph_run_properties(paragraph)
    for child in list(paragraph._p):
        if child.tag == qn("w:r"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if template_rpr is not None:
        run._r.insert(0, deepcopy(template_rpr))
        set_run_font(run, size=run.font.size.pt if run.font.size else None, bold=bool(run.bold) if run.bold is not None else None)
    else:
        set_run_font(run, size=None, bold=None)


def first_paragraph_run_properties(paragraph):
    for run in paragraph.runs:
        if run._r.rPr is not None:
            return deepcopy(run._r.rPr)
    return None


def style_table(table):
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    set_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=20, start=45, bottom=20, end=45)


def set_cell(
    cell,
    text,
    bold=False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    size=None,
    line_spacing=1,
    empty_text=EMPTY,
):
    template_rpr = first_cell_run_properties(cell)
    is_empty = text is None or text == ""
    rendered_text = empty_text if is_empty else normalize_dash_text(str(text))
    if rendered_text == EMPTY:
        align = WD_ALIGN_PARAGRAPH.CENTER
    cell.text = ""
    lines = rendered_text.splitlines() or [""]
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = line_spacing
        run = paragraph.add_run(line)
        if template_rpr is not None and size is None:
            run._r.insert(0, deepcopy(template_rpr))
        set_run_font(run, size=size, bold=bold)


def first_cell_run_properties(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run._r.rPr is not None:
                return deepcopy(run._r.rPr)
    return None


def normalize_dash_text(text):
    return EMPTY if text.strip() in {"-", "－", "–", "—"} else text


def set_sub_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    set_cell(cell, text, bold=bold, align=align, size=COMPACT_TABLE_FONT_SIZE, line_spacing=0.86)


def merge_write(table, top, left, bottom, right, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=None):
    cell = table.cell(top, left).merge(table.cell(bottom, right))
    set_cell(cell, text, bold=bold, align=align, size=size, line_spacing=0.86 if size else 1)
    return cell


def set_run_font(run, size=9, bold=False):
    if bold is not None:
        run.bold = bold
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT_NAME)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)


def set_table_borders(table):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_grid(table, widths):
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(index, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")


def value(raw):
    if raw is None or raw == "":
        return EMPTY
    return str(raw)


if __name__ == "__main__":
    main()
