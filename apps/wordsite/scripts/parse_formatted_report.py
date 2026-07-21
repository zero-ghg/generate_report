import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

EMPTY_VALUES = {"", "-", "－", "–", "—", "—"}
CHAPTER_PATTERN = re.compile(r"^[一二三四五六七八九十百]+、")
COVER_LABELS = {
    "委托单位名称": "clientName",
    "受检项目名称": "projectName",
    "受检项目地址": "projectAddress",
    "联   系   人": "contactName",
    "电        话": "contactPhone",
    "本次检测时间": "inspectedDate",
    "下次检测时间": "nextDate",
    "检测机构名称": "agencyName",
    "检测机构地址": "agencyAddress",
    "检测机构电话": "agencyPhone",
}
SECTION_KEYWORDS = {
    "overview": "子项目表（概况）",
    "power": "子项目表（低压电源系统）",
    "electronic": "子项目表（电子信息系统）",
    "grounding": "子项目表（接地电阻）",
    "transition": "子项目表（过渡电阻）",
    "spd": "子项目表（SPD明细表）",
    "spdTest": "子项目表（SPD测试表）",
    "legend": "现场平面示意图",
}


def parse_formatted_report_docx(source):
    if isinstance(source, (str, Path)):
        document = Document(str(source))
    else:
        document = Document(BytesIO(source))

    blocks = list(iter_body_blocks(document))
    toc_seen = False
    current_section = None
    summary_table = None
    overview_projects = []
    power_projects = []
    electronic_projects = []
    report_tables = {
        "grounding": [],
        "transition": [],
        "spd": [],
        "spdTest": [],
    }
    legend = {}
    current_subproject = None
    current_measurement_kind = None
    current_place_name = None
    # Word list numbering is often continued across several measurement
    # tables.  Keep the counter for the entire document instead of restarting
    # at D1/S1 for each page-sized table.
    measurement_numbering_counters = {}

    for kind, item in blocks:
        if kind == "sdt":
            toc_seen = True
            continue

        if kind == "p":
            text = normalize_text(item.text)
            if not text:
                continue

            if not toc_seen:
                continue

            section = detect_section(text)
            if section == "summary":
                current_section = "summary"
                current_subproject = None
                current_measurement_kind = None
                continue
            if section == "certificate":
                current_section = "certificate"
                continue
            if section:
                current_section = section
                current_subproject = None
                current_measurement_kind = section if section in report_tables else None
                current_place_name = None
                if section == "legend":
                    legend = {}
                continue

            if current_section == "legend" and text == "检测点位":
                continue

        if kind != "tbl":
            continue

        table = item
        if current_section == "summary" and summary_table is None:
            summary_table = table
            continue

        if current_section in {"overview", "power", "electronic"}:
            project = parse_subproject_table(table)
            if project:
                target = {
                    "overview": overview_projects,
                    "power": power_projects,
                    "electronic": electronic_projects,
                }[current_section]
                if target and target[-1].get("projectName") == project.get("projectName"):
                    target[-1]["rows"].extend(project.get("rows", []))
                else:
                    target.append(project)
            continue

        if current_section in report_tables:
            kind = "spd_test" if current_section == "spdTest" else current_section
            place_name = parse_measurement_place_name(table, kind)
            rows = parse_measurement_table(table, kind, measurement_numbering_counters)
            for row in rows:
                row["placeName"] = place_name or current_place_name or ""
            if place_name:
                current_place_name = place_name
            report_tables[current_section].extend(rows)
            continue

        if current_section == "legend" and is_test_point_table(table):
            legend["testPoints"] = parse_test_point_table(table)

    general = parse_summary_table(summary_table) if summary_table is not None else {}
    cover_general = parse_cover_paragraphs(document, toc_seen)
    for key, value in cover_general.items():
        if has_value(value) and not has_value(general.get(key)):
            general[key] = value

    return {
        "assistant": {
            "general": general,
            "overview": {"projects": overview_projects},
            "power": {"projects": power_projects},
            "electronic": {"projects": electronic_projects},
        },
        "reportTables": report_tables,
        "legend": legend,
    }


def iter_body_blocks(document):
    for child in document.element.body:
        if child.tag == qn("w:p"):
            yield "p", Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield "tbl", Table(child, document)
        elif child.tag == qn("w:sdt"):
            yield "sdt", child


def detect_section(text):
    if "总表" in text and CHAPTER_PATTERN.match(text):
        return "summary"
    if "资质证书" in text and CHAPTER_PATTERN.match(text):
        return "certificate"
    for key, keyword in SECTION_KEYWORDS.items():
        if keyword in text and CHAPTER_PATTERN.match(text):
            return key
    return None


def parse_cover_paragraphs(document, toc_seen):
    general = {}
    statement_lines = []
    in_statement = False
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        if "目录" in text.replace(" ", "") or CHAPTER_PATTERN.match(text):
            break
        compact_text = text.replace(" ", "")
        if compact_text == "声明":
            in_statement = True
            continue
        if in_statement:
            statement_lines.append(text)
        for label, field_name in COVER_LABELS.items():
            if not text.startswith(label):
                continue
            value = normalize_text(text[len(label) :])
            value = re.sub(r"（盖章）.*$", "", value).strip()
            if has_value(value):
                general[field_name] = value
        if text.startswith("报告编号"):
            general["reportNumber"] = normalize_text(text.split("：", 1)[-1])
        elif text.startswith("（第") and "册" in text:
            general["reportBookInfo"] = text
        elif text == "雷电防护装置检测报告":
            general["coverTitle"] = text
        elif text.lower().startswith("inspection report"):
            general["coverTitleEn"] = text
    if statement_lines:
        general["statementContent"] = "\n".join(statement_lines)
    return general


def parse_summary_table(table):
    conclusion_row = find_summary_row(table, "检测结论")
    if conclusion_row is not None:
        return parse_summary_table_compact(table)
    return parse_summary_table_full(table)


def parse_summary_table_compact(table):
    general = {
        "clientName": summary_value_after_label(table, 0, "委托单位名称", 3),
        "projectName": summary_value_after_label(table, 1, "受检项目名称", 3),
        "projectAddress": summary_value_after_label(table, 2, "受检项目地址", 3),
        "contactDepartment": summary_value_after_label(table, 3, "联系部门", 3),
        "contactName": summary_value_after_label(table, 3, "联系人", 7),
        "contactPhone": summary_value_after_label(table, 3, "联系电话", 15),
        "inspectedDate": summary_value_after_label(table, 4, "本次检测日期", 3),
        "nextDate": summary_value_after_label(table, 5, "下次检测日期", 3),
        "detectionType": parse_detection_type(summary_value_after_label(table, 5, "检测分类", 12)),
        "detectionBasis": summary_value_after_label(table, 6, "主要检测依据", 3),
        "equipment": [],
        "projectItems": [],
    }

    section = None
    equipment_columns = {"name": 1, "model": 4, "serial": 7, "range": 9, "calibrationDate": 12}
    project_columns = {"id": 0, "name": 2, "lightningCategory": 6, "lightningProtectionLevel": 11, "pages": 14}
    for row_index in range(7, len(table.rows)):
        row_values = summary_row_values(table, row_index)
        row_label = cell_text(table, row_index, 0)
        row_label_alt = cell_text(table, row_index, 1)
        joined_label = "".join(row_values)

        if is_equipment_header_row(table, row_index):
            equipment_columns = summary_header_columns(
                table,
                row_index,
                {
                    "name": ("仪器名称",),
                    "model": ("型号",),
                    "serial": ("编号",),
                    "range": ("测量范围",),
                    "calibrationDate": (
                        "检定/校准有效截止日期",
                        "鉴定/校准有效截止日期",
                        "校准日期",
                        "有效截止日期",
                    ),
                },
                equipment_columns,
            )
            section = "equipment_header"
            continue
        if "检测子项目" in joined_label or "检测项目列表" in joined_label or "子项目列表" in joined_label:
            project_columns = summary_header_columns(
                table,
                row_index,
                {
                    "id": ("序号", "编号"),
                    "name": ("检测子项目名称（场所）", "检测子项目名称", "子项目名称（场所）", "子项目名称"),
                    "lightningCategory": ("防雷类别",),
                    "lightningProtectionLevel": ("雷电防护等级",),
                    "pages": ("页码",),
                },
                project_columns,
            )
            section = "project_header"
            continue
        if row_label in {"序号", "编号"} or row_label_alt in {"检测子项目名称", "子项目名称"}:
            project_columns = summary_header_columns(
                table,
                row_index,
                {
                    "id": ("序号", "编号"),
                    "name": ("检测子项目名称（场所）", "检测子项目名称", "子项目名称（场所）", "子项目名称"),
                    "lightningCategory": ("防雷类别",),
                    "lightningProtectionLevel": ("雷电防护等级",),
                    "pages": ("页码",),
                },
                project_columns,
            )
            section = "project_header"
            continue
        if "检测结论" in joined_label or joined_label.startswith("经过") or "得出如下结论" in joined_label:
            general["conclusion"] = parse_conclusion_row(table, row_index)
            section = "conclusion"
            continue
        if row_label == "检测人" or row_label_alt == "检测人":
            general["tester"] = summary_value_after_label(table, row_index, "检测人", 2)
            general["reviewer"] = summary_value_after_label(table, row_index, "校核人", 8)
            general["signer"] = summary_value_after_label(table, row_index, "授权签字人", 13)
            break
        if section == "conclusion":
            continue

        if section == "equipment_header":
            item = {
                key: cell_text(table, row_index, column)
                for key, column in equipment_columns.items()
            }
            if item["name"] in {"检测项目列表", "检测子项目列表"}:
                section = "project_header"
                continue
            if any(has_value(value) for value in item.values()):
                general["equipment"].append(item)
            continue

        if section == "project_header":
            item = {
                key: cell_text(table, row_index, column)
                for key, column in project_columns.items()
            }
            if any(has_value(value) for value in item.values()):
                general["projectItems"].append(item)

    return general


def parse_summary_table_full(table):
    general = {
        "clientName": summary_value_after_label(table, 0, "委托单位名称", 3),
        "projectName": summary_value_after_label(table, 1, "受检项目名称", 3),
        "projectAddress": summary_value_after_label(table, 2, "受检项目地址", 3),
        "contactDepartment": summary_value_after_label(table, 3, "联系部门", 3),
        "contactName": summary_value_after_label(table, 3, "联系人", 6),
        "contactPhone": summary_value_after_label(table, 3, "联系电话", 14),
        "inspectedDate": summary_value_after_label(table, 4, "本次检测日期", 3),
        "nextDate": summary_value_after_label(table, 5, "下次检测日期", 3),
        "detectionType": parse_detection_type(summary_value_after_label(table, 5, "检测分类", 11)),
        "detectionBasis": summary_value_after_label(table, 6, "主要检测依据", 3),
        "equipment": [],
        "projectItems": [],
    }

    for row_index in range(8, min(14, len(table.rows))):
        item = {
            "name": cell_text(table, row_index, 1),
            "model": cell_text(table, row_index, 4),
            "serial": cell_text(table, row_index, 7),
            "range": cell_text(table, row_index, 10),
            "calibrationDate": cell_text(table, row_index, 14),
        }
        if any(has_value(value) for value in item.values()):
            general["equipment"].append(item)

    for row_index in range(16, len(table.rows)):
        label = cell_text(table, row_index, 0)
        if "备注" in label or "检测人" in label:
            break
        item = {
            "id": cell_text(table, row_index, 0),
            "name": cell_text(table, row_index, 2),
            "lightningCategory": cell_text(table, row_index, 5),
            "lightningProtectionLevel": cell_text(table, row_index, 9),
            "pages": cell_text(table, row_index, 13),
        }
        if any(has_value(value) for value in item.values()):
            general["projectItems"].append(item)

    parse_summary_footer(table, general, compact=False)
    return general


def parse_summary_footer(table, general, compact):
    conclusion_row = 19 if compact else 27
    signature_row = 20 if compact else 28
    if len(table.rows) > conclusion_row:
        general["conclusion"] = parse_conclusion_cell(table, conclusion_row, 2)
    if len(table.rows) > signature_row:
        if compact:
            general["tester"] = cell_text(table, signature_row, 2)
            general["reviewer"] = cell_text(table, signature_row, 9)
            general["signer"] = cell_text(table, signature_row, 16)
        else:
            general["tester"] = cell_text(table, signature_row, 2)
            general["reviewer"] = cell_text(table, signature_row, 8)
            general["signer"] = cell_text(table, signature_row, 15)
    return general


def parse_conclusion_cell(table, row_index, col_index):
    try:
        cell = table.cell(row_index, col_index)
    except IndexError:
        return ""
    lines = [normalize_text(paragraph.text) for paragraph in cell.paragraphs]
    lines = [line for line in lines if line and "检测机构（检测专用章）" not in line and not line.startswith("签发日期")]
    return "\n".join(lines).strip()


def summary_row_values(table, row_index):
    values = []
    for col_index in range(len(table.columns)):
        value = cell_text(table, row_index, col_index)
        if value and (not values or values[-1] != value):
            values.append(value)
    return values


def summary_value_after_label(table, row_index, label, fallback_col_index):
    values = summary_row_values(table, row_index)
    compact_label = re.sub(r"\s+", "", label)
    for index, value in enumerate(values):
        compact_value = re.sub(r"\s+", "", value)
        if compact_value != compact_label:
            continue
        if index + 1 < len(values):
            next_value = values[index + 1]
            # An empty merged value (often rendered as "—") disappears during
            # normalization.  Do not then consume the next field label as data.
            field_labels = {
                "联系部门", "联系人", "联系电话", "检测人", "校核人", "授权签字人",
                "本次检测日期", "下次检测日期", "检测分类",
            }
            if re.sub(r"\s+", "", next_value) in field_labels:
                return ""
            return next_value
        break
    return cell_text(table, row_index, fallback_col_index)


def summary_header_columns(table, row_index, labels_by_key, fallback):
    columns = dict(fallback)
    for key, labels in labels_by_key.items():
        compact_labels = {re.sub(r"\s+", "", label) for label in labels}
        for col_index in range(len(table.columns)):
            value = re.sub(r"\s+", "", cell_text(table, row_index, col_index))
            if value in compact_labels:
                columns[key] = col_index
                break
    return columns


def find_summary_row(table, label):
    compact_label = re.sub(r"\s+", "", label)
    for row_index in range(len(table.rows)):
        if any(compact_label in re.sub(r"\s+", "", value) for value in summary_row_values(table, row_index)):
            return row_index
    return None


def parse_conclusion_row(table, row_index):
    values = summary_row_values(table, row_index)
    lines = []
    for value in values:
        compact_value = re.sub(r"\s+", "", value)
        if compact_value == "检测结论":
            continue
        value = re.sub(r"^检测结论[：:]?", "", value).strip()
        value = re.sub(r"检测机构（检测专用章）.*$", "", value).strip()
        value = re.sub(r"签发日期[：:]?.*$", "", value).strip()
        if value:
            lines.append(value)
    return "\n".join(lines).strip()


def parse_subproject_table(table):
    if len(table.rows) < 3:
        return None

    positions = subproject_positions(table)
    project = {
        "projectName": cell_text(table, 0, positions["name"]),
        "inspectionDate": cell_text(table, 0, positions["date"]),
        "rows": [],
    }
    if not has_value(project["projectName"]) and not has_value(project["inspectionDate"]):
        return None

    first_data_row = 2
    last_row_index = len(table.rows) - 1
    if is_remark_table_row(table.rows[-1]):
        last_row_index -= 1

    for row_index in range(first_data_row, last_row_index + 1):
        row = {
            "category": cell_text(table, row_index, positions["category"]) if positions["category"] is not None else "",
            "subcategory": cell_text(table, row_index, positions["subcategory"]) if positions["subcategory"] is not None else "",
            "content": cell_text(table, row_index, positions["content"]),
            "standard": cell_text(table, row_index, positions["standard"]),
            "result": cell_text(table, row_index, positions["result"]),
            "conclusion": cell_text(table, row_index, positions["conclusion"]),
        }
        if any(has_value(value) for value in row.values()):
            project["rows"].append(row)
    return project if project["rows"] else project if has_value(project["projectName"]) else None


def subproject_positions(table):
    if len(table.columns) == 7:
        return {
            "name": 2,
            "date": 5,
            "category": 0,
            "subcategory": None,
            "content": 1,
            "standard": 3,
            "result": 4,
            "conclusion": 6,
        }
    return {
        "name": 3,
        "date": 6,
        "category": 0,
        "subcategory": 1,
        "content": 2,
        "standard": 4,
        "result": 5,
        "conclusion": 7,
    }


def parse_measurement_place_name(table, kind):
    positions = measurement_positions(kind)
    return cell_text(table, 0, positions["name"])


def parse_measurement_table(table, kind, numbering_counters=None):
    if kind == "spd_test":
        return parse_spd_test_table(table)
    return parse_standard_measurement_table(table, kind, numbering_counters)


def parse_standard_measurement_table(table, kind, numbering_counters=None):
    rows = []
    numbering_counters = numbering_counters if numbering_counters is not None else {}
    first_data_row = 2
    last_row_index = len(table.rows) - 1
    if is_remark_table_row(table.rows[-1]):
        last_row_index -= 1

    for row_index in range(first_data_row, last_row_index + 1):
        if is_remark_table_row(table.rows[row_index]):
            continue
        if kind == "grounding":
            row = parse_grounding_row(table, row_index)
        elif kind == "transition":
            row = parse_transition_row(table, row_index)
        elif kind == "spd":
            row = parse_spd_row(table, row_index)
        else:
            row = {}
        if not measurement_row_has_payload(row):
            continue
        if row and not row.get("marker"):
            row["marker"] = automatic_cell_number(
                table.cell(row_index, 0), numbering_counters
            )
        rows.append(row)
    return rows


def measurement_row_has_payload(row):
    """Ignore empty template rows and merged remark rows before auto-numbering."""
    if not row:
        return False
    for key, value in row.items():
        if key in {"marker", "placeName"}:
            continue
        text = normalize_text(value)
        if not text or text in EMPTY_VALUES or text == "/" or text.startswith("备注"):
            continue
        return True
    return False


def automatic_cell_number(cell, counters):
    """Return the visible Word list number when it is not part of ``cell.text``."""
    for paragraph in cell.paragraphs:
        paragraph_properties = paragraph._p.pPr
        numbering_properties = (
            paragraph_properties.numPr if paragraph_properties is not None else None
        )
        if numbering_properties is None or numbering_properties.numId is None:
            continue

        num_id = int(numbering_properties.numId.val)
        level = (
            int(numbering_properties.ilvl.val)
            if numbering_properties.ilvl is not None
            else 0
        )
        definition = numbering_definition(paragraph, num_id, level)
        if definition is None:
            continue

        start, number_format, level_text = definition
        counter_key = (num_id, level)
        number = counters.get(counter_key, start)
        counters[counter_key] = number + 1
        return format_automatic_number(number, number_format, level_text, level)
    return ""


def numbering_definition(paragraph, num_id, level):
    numbering_root = paragraph.part.numbering_part.element
    num_element = next(
        (
            item
            for item in numbering_root.findall(qn("w:num"))
            if int(item.get(qn("w:numId"))) == num_id
        ),
        None,
    )
    if num_element is None:
        return None

    abstract_id_element = num_element.find(qn("w:abstractNumId"))
    if abstract_id_element is None:
        return None
    abstract_id = int(abstract_id_element.get(qn("w:val")))
    abstract_element = next(
        (
            item
            for item in numbering_root.findall(qn("w:abstractNum"))
            if int(item.get(qn("w:abstractNumId"))) == abstract_id
        ),
        None,
    )
    if abstract_element is None:
        return None

    level_element = next(
        (
            item
            for item in abstract_element.findall(qn("w:lvl"))
            if int(item.get(qn("w:ilvl"))) == level
        ),
        None,
    )
    if level_element is None:
        return None

    start = 1
    start_element = level_element.find(qn("w:start"))
    if start_element is not None:
        start = int(start_element.get(qn("w:val")))

    for override in num_element.findall(qn("w:lvlOverride")):
        if int(override.get(qn("w:ilvl"))) != level:
            continue
        start_override = override.find(qn("w:startOverride"))
        if start_override is not None:
            start = int(start_override.get(qn("w:val")))
        break

    format_element = level_element.find(qn("w:numFmt"))
    text_element = level_element.find(qn("w:lvlText"))
    number_format = (
        format_element.get(qn("w:val")) if format_element is not None else "decimal"
    )
    level_text = text_element.get(qn("w:val")) if text_element is not None else "%1"
    return start, number_format, level_text


def format_automatic_number(number, number_format, level_text, level):
    if number_format == "upperLetter":
        rendered_number = alphabetic_number(number).upper()
    elif number_format == "lowerLetter":
        rendered_number = alphabetic_number(number).lower()
    elif number_format == "decimalZero":
        rendered_number = f"{number:02d}"
    else:
        rendered_number = str(number)
    return level_text.replace(f"%{level + 1}", rendered_number)


def alphabetic_number(number):
    result = ""
    while number > 0:
        number, remainder = divmod(number - 1, 26)
        result = chr(ord("A") + remainder) + result
    return result


def measurement_column_positions(table, kind):
    """Resolve columns from each table's real header and merged-cell grid."""
    positions = {}
    seen_cells = set()
    header_row = table.rows[1] if len(table.rows) > 1 else table.rows[0]
    labels = {
        "marker": ("\u7f16\u53f7",),
        "location": ("\u6240\u5728\u4f4d\u7f6e",),
        "conductor": ("\u8fde\u63a5\u5bfc\u4f53",),
        "protection": ("\u9632\u96f7\u5206\u533a",),
        "standard": ("\u6807\u51c6\u503c",),
        "measured": ("\u6d4b\u8bd5\u503c", "\u5b9e\u6d4b\u503c"),
        "result": ("\u7ed3\u8bba",),
        "spd_model": ("SPD\u578b\u53f7",),
        "install_location": ("\u5b89\u88c5\u4f4d\u7f6e",),
        "wire_length": ("\u63a5\u7ebf\u957f\u5ea6",),
        "spd_level": ("SPD\u7ea7\u522b",),
        "install_quantity": ("\u5b89\u88c5\u6570\u91cf",),
    }
    for column, cell in enumerate(header_row.cells):
        cell_key = id(cell._tc)
        if cell_key in seen_cells:
            continue
        seen_cells.add(cell_key)
        header = normalize_text(cell.text).replace("\n", "").replace(" ", "")
        for field, candidates in labels.items():
            if field not in positions and any(candidate in header for candidate in candidates):
                positions[field] = column

    defaults = {
        "grounding": {"marker": 0, "location": 1, "conductor": 3, "protection": 4, "standard": 5, "measured": 6, "result": 7},
        "transition": {"marker": 0, "location": 1, "conductor": 3, "protection": 4, "standard": 5, "measured": 6, "result": 7},
        "spd": {"marker": 0, "spd_model": 1, "install_location": 2, "wire_length": 3, "spd_level": 4, "install_quantity": 5, "measured": 6, "result": 7},
    }[kind]
    return {**defaults, **positions}


def parse_grounding_row(table, row_index):
    positions = measurement_column_positions(table, "grounding")
    location_text = cell_text(table, row_index, positions["location"])
    work_location, equipment_name = split_location_text(location_text)
    return {
        "marker": cell_text(table, row_index, positions["marker"]),
        "workLocation": work_location,
        "equipmentName": equipment_name,
        "conductorSpec": cell_placeholder_text(table, row_index, positions["conductor"]),
        "protectionZone": cell_text(table, row_index, positions["protection"]),
        "standardValue": cell_text(table, row_index, positions["standard"]),
        "measuredValue": cell_text(table, row_index, positions["measured"]),
        "result": cell_text(table, row_index, positions["result"]),
    }


def parse_transition_row(table, row_index):
    positions = measurement_column_positions(table, "transition")
    location_text = cell_text(table, row_index, positions["location"])
    parts = split_joined_text(location_text)
    # 合并单元格依次表示所在位置、设备名称、基准点，字段之间不能互相借值。
    reference_point = parts[2] if len(parts) > 2 else ""
    return {
        "marker": cell_text(table, row_index, positions["marker"]),
        "workLocation": parts[0] if len(parts) > 0 else "",
        "equipmentName": parts[1] if len(parts) > 1 else "",
        "referencePoint": reference_point,
        "conductorSpec": cell_placeholder_text(table, row_index, positions["conductor"]),
        "protectionZone": cell_text(table, row_index, positions["protection"]),
        "standardValue": cell_text(table, row_index, positions["standard"]),
        "measuredValue": cell_text(table, row_index, positions["measured"]),
        "result": cell_text(table, row_index, positions["result"]),
    }


def parse_spd_row(table, row_index):
    positions = measurement_column_positions(table, "spd")
    return {
        "marker": cell_text(table, row_index, positions["marker"]),
        "spdModel": cell_text(table, row_index, positions["spd_model"]),
        "installLocation": cell_text(table, row_index, positions["install_location"]),
        "wireLength": cell_text(table, row_index, positions["wire_length"]),
        "spdLevel": cell_text(table, row_index, positions["spd_level"]),
        "installQuantity": cell_text(table, row_index, positions["install_quantity"]),
        "measuredValue": cell_text(table, row_index, positions["measured"]),
        "result": cell_text(table, row_index, positions["result"]),
    }


def parse_spd_test_table(table):
    rows = []
    first_data_row = 4
    last_row_index = len(table.rows) - 1
    if is_remark_table_row(table.rows[-1]):
        last_row_index -= 1

    row_index = first_data_row
    while row_index < last_row_index:
        row = {
            "marker": cell_text(table, row_index, 0),
            "spdModel": cell_text(table, row_index, 1),
            "installLocation": cell_text(table, row_index, 2),
            "u1maL1": cell_text(table, row_index, 3),
            "u1maL2": cell_text(table, row_index, 4),
            "leakageL1": cell_text(table, row_index, 5),
            "leakageL2": cell_text(table, row_index, 6),
            "insulationL1": cell_text(table, row_index, 7),
            "insulationL2": cell_text(table, row_index, 8),
            "result": cell_text(table, row_index, 9),
        }
        next_row_index = row_index + 1
        if next_row_index <= last_row_index:
            row.update(
                {
                    "u1maL3": cell_text(table, next_row_index, 3),
                    "u1maN": cell_text(table, next_row_index, 4),
                    "leakageL3": cell_text(table, next_row_index, 5),
                    "leakageN": cell_text(table, next_row_index, 6),
                    "insulationL3": cell_text(table, next_row_index, 7),
                    "insulationN": cell_text(table, next_row_index, 8),
                }
            )
        if any(has_value(value) for key, value in row.items() if key != "result"):
            rows.append(row)
        row_index += 2
    return rows


def measurement_positions(kind):
    if kind == "spd_test":
        return {"name": 2, "date": 7}
    return {"name": 2, "date": 6}


def is_test_point_table(table):
    if not table.rows:
        return False
    header = "".join(cell_text(table, 0, col_index) for col_index in range(len(table.columns)))
    return "点位编号" in header or "检测类型" in header


def parse_test_point_table(table):
    columns = []
    for col_index in range(len(table.columns)):
        header = cell_text(table, 0, col_index)
        field = {
            "点位编号": "label",
            "检测类型": "reportType",
            "设备名称": "equipmentName",
            "位置": "workLocation",
            "标准值": "standardValue",
            "实测值": "measuredValue",
            "结论": "result",
        }.get(header)
        if field:
            columns.append((col_index, field))

    points = []
    for row_index in range(1, len(table.rows)):
        fields = {}
        item = {"reportFields": fields}
        for col_index, field in columns:
            value = cell_text(table, row_index, col_index)
            if field == "label":
                item["label"] = value
            elif field == "reportType":
                item["reportType"] = value
            else:
                fields[field] = value
        if any(has_value(value) for value in [item.get("label"), item.get("reportType"), *fields.values()]):
            points.append(item)
    return points


def parse_detection_type(text):
    if not has_value(text):
        return text
    if "☑定期检测" in text:
        return "定期检测"
    if "☑验收检测" in text:
        return "验收检测"
    return text


def split_location_text(text):
    parts = split_joined_text(text)
    return parts[0] if len(parts) > 0 else "", parts[1] if len(parts) > 1 else ""


def split_joined_text(text):
    if not has_value(text):
        return []
    return [part.strip() for part in re.split(r"[，,]", str(text)) if part.strip()]


def is_equipment_header_row(table, row_index):
    row_label = cell_text(table, row_index, 0)
    row_label_alt = cell_text(table, row_index, 1)
    joined_label = f"{row_label}{row_label_alt}"
    if "主要" not in joined_label or "设备" not in joined_label:
        return False
    header_names = {"仪器名称", "型号", "编号", "测量范围", "校准日期"}
    return row_label_alt in header_names or cell_text(table, row_index, 2) in header_names


def is_remark_table_row(row):
    if not row.cells:
        return False
    return row.cells[0].text.strip().startswith("备注")


def cell_text(table, row_index, col_index):
    if row_index < 0 or col_index < 0 or row_index >= len(table.rows) or col_index >= len(table.columns):
        return ""
    try:
        return clean_value(table.cell(row_index, col_index).text)
    except IndexError:
        return ""


def cell_placeholder_text(table, row_index, col_index):
    """保留报告表格中明确填写的横线占位符，避免前端误判为缺少必填项。"""
    try:
        value = normalize_text(table.cell(row_index, col_index).text)
    except IndexError:
        return ""
    if value in EMPTY_VALUES:
        return "—" if value else ""
    return value


def clean_value(text):
    value = normalize_text(text)
    if value in EMPTY_VALUES:
        return ""
    return value


def normalize_text(text):
    if text is None:
        return ""
    return re.sub(r"\s+", " ", str(text)).strip()


def has_value(text):
    return clean_value(text) != ""
