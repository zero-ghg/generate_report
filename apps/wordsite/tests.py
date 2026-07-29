from io import StringIO

import ezdxf
from django.test import SimpleTestCase
from docx import Document
from docx.shared import RGBColor

from apps.wordsite.scripts.generate_formatted_report import (
    expand_compact_subproject_category_column,
    extract_legend_canvases,
    merge_subproject_category_columns,
    measurement_group_capacity,
    repeat_measurement_header_rows,
    trim_measurement_rows_to_fit,
)
from apps.wordsite.scripts.parse_formatted_report import parse_grounding_row, parse_subproject_table, parse_summary_table, parse_transition_row
from apps.wordsite.scripts.parse_dwg_workspace import (
    _assign_interaction_groups,
    _cad_glyph_outlines_complete,
    _collapse_paired_flange_test_points,
    _fit_cad_text_to_enclosing_boxes,
    _is_orphan_legend_swatch,
    _match_marker_text,
    _marker_text_visual_box,
    _merge_legend_hatch_parts,
    _report_fields,
    _remove_metal_roof_legend_empty_frame,
    _remove_redundant_large_hatch_outlines,
    _report_marker_for_display_label,
    _anchor_room_labels_to_enclosing_boxes,
    _separate_connector_identifier_labels,
    _separate_overlapping_marker_label,
    _triangle_side_from_path,
    _wrapped_cad_text,
    finalize_existing_report_workspace,
    parse_dwg_workspace,
)


# DWG/DXF 工作区解析测试：验证图纸文字与检测表记录共享 id，并保留未匹配记录。
class DwgWorkspaceParserTests(SimpleTestCase):
    def test_exports_each_child_drawing_as_a_separate_legend_canvas(self):
        legend = {
            "boardHeight": 600,
            "boardWidth": 900,
            "drawingGroups": [
                {"height": 300, "id": "drawing-a", "name": "A图", "width": 400},
                {"height": 320, "id": "drawing-b", "name": "B图", "width": 500},
            ],
            "paths": [
                {"drawingId": "drawing-a", "points": [{"x": 10, "y": 10}]},
                {"drawingId": "drawing-b", "points": [{"x": 20, "y": 20}]},
            ],
        }

        canvases = extract_legend_canvases(legend)

        self.assertEqual(len(canvases), 2)
        self.assertEqual(canvases[0]["boardWidth"], 400)
        self.assertEqual(canvases[1]["boardHeight"], 320)
        self.assertEqual(canvases[0]["paths"][0]["drawingId"], "drawing-a")
        self.assertEqual(canvases[1]["paths"][0]["drawingId"], "drawing-b")

    def test_compact_overview_category_does_not_become_its_own_subcategory(self):
        document = Document()
        table = document.add_table(rows=3, cols=8)
        table.cell(0, 3).text = "办公区"
        table.cell(0, 6).text = "2025年3月7日"
        table.cell(2, 0).text = "引 下 线"
        table.cell(2, 1).text = ""
        table.cell(2, 2).text = "对应接闪器"
        table.cell(2, 4).text = "接闪带（网）"
        table.cell(2, 5).text = "接闪带"
        table.cell(2, 6).text = "符合"
        table.cell(2, 0).merge(table.cell(2, 1))

        project = parse_subproject_table(table)

        self.assertEqual(project["rows"][0]["category"], "引 下 线")
        self.assertEqual(project["rows"][0]["subcategory"], "")

    def test_subproject_table_preserves_dash_placeholders(self):
        document = Document()
        table = document.add_table(rows=3, cols=8)
        table.cell(0, 3).text = "辅助用房"
        table.cell(0, 6).text = "2025年3月7日"
        table.cell(2, 0).text = "低压电源系统"
        table.cell(2, 2).text = "SPD安装位置"
        table.cell(2, 4).text = "-"
        table.cell(2, 5).text = "—"
        table.cell(2, 7).text = "－"

        project = parse_subproject_table(table)

        self.assertEqual(project["rows"][0]["standard"], "—")
        self.assertEqual(project["rows"][0]["result"], "—")
        self.assertEqual(project["rows"][0]["conclusion"], "—")

    def test_subproject_table_preserves_red_field_markers(self):
        document = Document()
        table = document.add_table(rows=3, cols=8)
        table.cell(0, 3).text = "辅助用房"
        table.cell(0, 6).text = "2025年3月7日"
        table.cell(2, 0).text = "低压电源系统"
        table.cell(2, 2).text = "SPD外部脱离器"
        red_cell = table.cell(2, 4)
        red_cell.text = ""
        run = red_cell.paragraphs[0].add_run("GB/T 21431-2023 5.5.6.7")
        run.font.color.rgb = RGBColor(255, 0, 0)
        table.cell(2, 5).text = "有/正常"
        table.cell(2, 7).text = "符合"

        project = parse_subproject_table(table)

        self.assertEqual(project["rows"][0]["standard"], "GB/T 21431-2023 5.5.6.7")
        self.assertEqual(project["rows"][0]["fieldColors"]["standard"], "#ff0000")

    def test_subproject_table_preserves_partial_red_runs(self):
        document = Document()
        table = document.add_table(rows=3, cols=8)
        table.cell(0, 3).text = "辅助用房"
        table.cell(0, 6).text = "2025年3月7日"
        table.cell(2, 0).text = "低压电源系统"
        table.cell(2, 2).text = "SPD外观"
        result_cell = table.cell(2, 5)
        result_cell.text = ""
        result_cell.paragraphs[0].add_run("表面平整，")
        red_run = result_cell.paragraphs[0].add_run("无划伤")
        red_run.font.color.rgb = RGBColor(255, 0, 0)
        result_cell.paragraphs[0].add_run("。")
        table.cell(2, 7).text = "符合"

        project = parse_subproject_table(table)

        self.assertEqual(project["rows"][0]["result"], "表面平整，无划伤。")
        self.assertEqual(
            project["rows"][0]["formattedFields"]["result"],
            [{"text": "表面平整，"}, {"text": "无划伤", "color": "#ff0000"}, {"text": "。"}],
        )

    def test_uses_triangle_geometry_for_left_facing_marker(self):
        # The legend label sits to the right of this marker, but the imported
        # solid triangle itself is what defines its direction.
        path = {
            "points": [
                {"x": 10, "y": 0},
                {"x": 0, "y": 5},
                {"x": 10, "y": 10},
                {"x": 10, "y": 0},
            ],
        }

        self.assertEqual(_triangle_side_from_path(path), "left")

    def test_moves_sc_connector_identifier_above_its_box(self):
        texts = [{
            "fontSize": 10,
            "height": 12,
            "text": "SC1",
            "width": 20,
            "x": 110,
            "y": 90,
        }]
        paths = [{
            "closed": True,
            "points": [
                {"x": 105, "y": 100},
                {"x": 135, "y": 100},
                {"x": 135, "y": 116},
                {"x": 105, "y": 116},
            ],
        }]

        _separate_connector_identifier_labels(texts, paths)

        self.assertLessEqual(texts[0]["y"] + texts[0]["height"], 97.5)

    def test_anchors_distribution_room_label_inside_lower_left_corner(self):
        texts = [{
            "fontSize": 14,
            "height": 18,
            "text": "配电室",
            "width": 56,
            "x": 180,
            "y": 180,
        }]
        def rectangle(left, top, right, bottom):
            return {
                "closed": True,
                "points": [
                    {"x": left, "y": top},
                    {"x": right, "y": top},
                    {"x": right, "y": bottom},
                    {"x": left, "y": bottom},
                ],
            }

        paths = [
            rectangle(100, 100, 300, 300),  # room frame
            rectangle(170, 170, 250, 210),  # misleading nearby equipment box
            rectangle(170, 115, 250, 145),
            rectangle(170, 220, 250, 250),
        ]

        _anchor_room_labels_to_enclosing_boxes(texts, paths)

        self.assertAlmostEqual(texts[0]["x"], 104)
        self.assertAlmostEqual(texts[0]["y"], 276)

    def test_keeps_short_cad_identifiers_on_one_line(self):
        item = {
            "cadBoxWidth": 5.5,
            "fontSize": 4,
            "fontWeight": 400,
            "text": "SC1",
            "widthFactor": 0.8,
        }

        self.assertEqual(_wrapped_cad_text(item), "SC1")

    def test_keeps_zero_padded_numeric_marker_on_one_line(self):
        item = {
            "cadBoxWidth": 5.5,
            "fontSize": 4,
            "fontWeight": 400,
            "text": "025",
            "widthFactor": 0.8,
        }

        self.assertEqual(_wrapped_cad_text(item), "025")

    def test_rejects_partial_cad_glyph_outlines(self):
        item = {
            "fontSize": 5.2,
            "text": "PLC柜",
            "widthFactor": 0.8,
        }
        only_letter_c = [[
            {"x": 0, "y": 0},
            {"x": 3.2, "y": 0},
            {"x": 3.2, "y": 5},
            {"x": 0, "y": 5},
        ]]
        complete_label = [[
            {"x": 0, "y": 0},
            {"x": 15, "y": 0},
            {"x": 15, "y": 5},
            {"x": 0, "y": 5},
        ]]

        self.assertFalse(_cad_glyph_outlines_complete(item, only_letter_c))
        self.assertTrue(_cad_glyph_outlines_complete(item, complete_label))

    def test_pads_measurement_rows_to_full_page_capacity(self):
        document = Document()
        table = document.add_table(rows=34, cols=8)
        table.cell(33, 0).text = "备注："

        trim_measurement_rows_to_fit(table, data_row_count=9, max_data_rows=30, kind="transition")

        self.assertEqual(len(table.rows), 34)
        self.assertEqual(table.rows[-2].cells[0].text, "")
        self.assertEqual(table.rows[-1].cells[0].text, "备注：")

    def test_pads_spd_test_rows_in_two_row_pairs(self):
        document = Document()
        table = document.add_table(rows=34, cols=10)
        table.cell(33, 0).text = "备注："

        trim_measurement_rows_to_fit(table, data_row_count=3, max_data_rows=15, kind="spd_test")

        self.assertEqual(len(table.rows), 35)
        self.assertEqual(table.rows[-2].cells[0].text, "")
        self.assertEqual(table.rows[-1].cells[0].text, "备注：")

    def test_measurement_capacity_keeps_one_table_for_multiple_pages(self):
        self.assertEqual(measurement_group_capacity(39, 30), 39)
        self.assertEqual(measurement_group_capacity(16, 15), 16)
        self.assertEqual(measurement_group_capacity(0, 30), 30)

    def test_multi_page_measurement_table_does_not_add_empty_data_rows(self):
        document = Document()
        table = document.add_table(rows=63, cols=8)
        table.cell(62, 0).text = "备注："

        trim_measurement_rows_to_fit(table, data_row_count=39, max_data_rows=39, kind="transition")

        self.assertEqual(len(table.rows), 43)
        self.assertEqual(table.rows[-2].cells[0].text, "")
        self.assertEqual(table.rows[-1].cells[0].text, "备注：")

    def test_measurement_headers_repeat_when_one_table_crosses_pages(self):
        document = Document()
        table = document.add_table(rows=5, cols=2)

        repeat_measurement_header_rows(table, "transition")

        self.assertEqual(len(table.rows[0]._tr.xpath("./w:trPr/w:tblHeader")), 1)
        self.assertEqual(len(table.rows[1]._tr.xpath("./w:trPr/w:tblHeader")), 1)
        self.assertEqual(len(table.rows[2]._tr.xpath("./w:trPr/w:tblHeader")), 0)

    def test_expands_compact_overview_table_for_subcategory_data(self):
        document = Document()
        table = document.add_table(rows=4, cols=7)
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(1, 1).merge(table.cell(1, 2))
        rows = [
            {"category": "接闪器", "subcategory": "类型"},
            {"category": "基本信息", "subcategory": "基本信息"},
        ]

        expand_compact_subproject_category_column(table, rows)

        self.assertEqual(len(table.columns), 8)
        self.assertIsNot(table.cell(2, 0)._tc, table.cell(2, 1)._tc)
        self.assertIs(table.cell(3, 0)._tc, table.cell(3, 1)._tc)

    def test_merges_repeated_overview_subcategory_cells(self):
        document = Document()
        table = document.add_table(rows=5, cols=8)
        positions = {"category": 0, "subcategory": 1}
        row_meta = [
            ("接闪器", "金属构件或金属屋面"),
            ("接闪器", "金属构件或金属屋面"),
            ("接闪器", "类型"),
        ]

        merge_subproject_category_columns(table, 2, row_meta, positions)

        self.assertIs(table.cell(2, 1)._tc, table.cell(3, 1)._tc)
        self.assertIsNot(table.cell(3, 1)._tc, table.cell(4, 1)._tc)

    def test_parses_expanded_compact_summary_without_treating_conclusion_as_project(self):
        document = Document()
        table = document.add_table(rows=23, cols=18)

        table.cell(0, 0).text = "委托单位名称"
        table.cell(0, 3).text = "天津北辉科技集团有限公司凤凰路加油站"
        table.cell(1, 0).text = "受检项目名称"
        table.cell(1, 3).text = "凤凰路加油站雷电防护装置检测项目"
        table.cell(2, 0).text = "受检项目地址"
        table.cell(2, 3).text = "天津市河东区凤凰路8号"
        table.cell(3, 0).text = "联系部门"
        table.cell(3, 2).text = "办公室"
        table.cell(3, 3).text = "联系人"
        table.cell(3, 7).text = "张志强"
        table.cell(3, 12).text = "联系电话"
        table.cell(3, 15).text = "18722562593"
        table.cell(4, 0).text = "本次检测日期"
        table.cell(4, 3).text = "2025年2月19日"
        table.cell(5, 0).text = "下次检测日期"
        table.cell(5, 3).text = "2025年8月20日前"
        table.cell(5, 6).text = "检测分类"
        table.cell(5, 12).text = "验收检测 定期检测"
        table.cell(6, 0).text = "主要检测依据"
        table.cell(6, 3).text = "检测规范"
        table.cell(7, 0).text = "主要检测设备"
        table.cell(7, 1).text = "仪器名称"
        table.cell(7, 3).text = "型号"
        table.cell(7, 4).text = "编号"
        table.cell(7, 9).text = "测量范围"
        table.cell(7, 14).text = "鉴定/校准有效截止日期"
        table.cell(8, 1).text = "接地电阻测试仪"
        table.cell(8, 3).text = "S-3019B"
        table.cell(8, 4).text = "052211679"
        table.cell(8, 9).text = "0.00Ω~2000Ω"
        table.cell(8, 14).text = "2024.12.2-2025.12.1"
        table.cell(14, 0).text = "检测项目列表"
        table.cell(15, 0).text = "序号"
        table.cell(15, 2).text = "检测子项目名称（场所）"
        table.cell(15, 8).text = "防雷类别"
        table.cell(15, 11).text = "雷电防护等级"
        table.cell(15, 16).text = "页码"
        table.cell(16, 0).text = "1"
        table.cell(16, 2).text = "加油区（金属罩棚）"
        table.cell(16, 8).text = "第二类"
        table.cell(16, 11).text = "—"
        table.cell(16, 16).text = "2、3、10、11"
        table.cell(21, 0).text = "检测结论"
        table.cell(21, 2).text = "经过现场检测，得出如下结论：报告合格。"
        table.cell(22, 0).text = "检测人"
        table.cell(22, 2).text = "检测员"

        general = parse_summary_table(table)

        self.assertEqual(general["contactDepartment"], "办公室")
        self.assertEqual(general["contactName"], "张志强")
        self.assertEqual(general["contactPhone"], "18722562593")
        self.assertEqual(general["conclusion"], "经过现场检测，得出如下结论：报告合格。")
        self.assertEqual([item["name"] for item in general["projectItems"]], ["加油区（金属罩棚）"])
        self.assertEqual(general["projectItems"][0]["lightningCategory"], "第二类")
        self.assertEqual(general["projectItems"][0]["lightningProtectionLevel"], "")
        self.assertEqual(general["projectItems"][0]["pages"], "2、3、10、11")
        self.assertEqual(general["equipment"][0]["model"], "S-3019B")
        self.assertEqual(general["equipment"][0]["serial"], "052211679")

    def test_parses_fourteen_column_summary_without_shifting_empty_signatures(self):
        document = Document()
        table = document.add_table(rows=28, cols=14)
        table.cell(0, 0).text = "委托单位名称"
        table.cell(0, 3).text = "测试单位"
        table.cell(1, 0).text = "受检项目名称"
        table.cell(1, 3).text = "测试项目"
        table.cell(2, 0).text = "受检项目地址"
        table.cell(2, 3).text = "测试地址"
        table.cell(3, 0).text = "联系部门"
        table.cell(3, 3).text = "—"
        table.cell(3, 4).text = "联系人"
        table.cell(3, 5).text = "王胜亚"
        table.cell(3, 9).text = "联系电话"
        table.cell(3, 12).text = "18222353738"
        table.cell(15, 0).text = "主要检测设备"
        table.cell(15, 1).text = "仪器名称"
        table.cell(15, 4).text = "型号"
        table.cell(15, 7).text = "编号"
        table.cell(15, 9).text = "测量范围"
        table.cell(15, 12).text = "检定/校准有效截止日期"
        table.cell(16, 1).text = "接地电阻测试仪"
        table.cell(16, 4).text = "S-3019B"
        table.cell(16, 7).text = "052211679"
        table.cell(16, 9).text = "0.00Ω~2000Ω"
        table.cell(16, 12).text = "2024.12.2-2025.12.1"
        table.cell(22, 0).text = "检测项目列表"
        table.cell(23, 0).text = "序号"
        table.cell(23, 2).text = "检测子项目名称（场所）"
        table.cell(23, 7).text = "防雷类别"
        table.cell(23, 9).text = "雷电防护等级"
        table.cell(23, 12).text = "页码"
        table.cell(24, 0).text = "1"
        table.cell(24, 2).text = "罩棚"
        table.cell(24, 7).text = "第二类"
        table.cell(26, 0).text = "检测结论"
        table.cell(26, 2).text = "检测合格"
        table.cell(27, 0).text = "检测人"
        table.cell(27, 6).text = "校核人"
        table.cell(27, 10).text = "授权签字人"

        general = parse_summary_table(table)

        self.assertEqual(general["contactDepartment"], "")
        self.assertEqual(general["contactName"], "王胜亚")
        self.assertEqual(general["contactPhone"], "18222353738")
        self.assertEqual(general["reviewer"], "")
        self.assertEqual(general["signer"], "")
        self.assertEqual(general["equipment"][0]["calibrationDate"], "2024.12.2-2025.12.1")
        self.assertEqual([item["name"] for item in general["projectItems"]], ["罩棚"])

    def test_moves_only_overlapping_marker_label_outside_symbol(self):
        marker_box = {"x": 96, "y": 96, "width": 8, "height": 8}
        overlapping = {
            "glyphBounds": {"x": 0, "y": 0, "width": 4, "height": 4},
            "x": 98,
            "y": 94,
        }
        clear = {
            "glyphBounds": {"x": 0, "y": 0, "width": 4, "height": 4},
            "x": 98,
            "y": 80,
        }

        _separate_overlapping_marker_label(overlapping, marker_box, -2, -6)
        _separate_overlapping_marker_label(clear, marker_box, -2, -20)

        overlapping_box = _marker_text_visual_box(overlapping)
        self.assertLessEqual(overlapping_box["y"] + overlapping_box["height"], 94.5)
        self.assertEqual((clear["x"], clear["y"]), (98, 80))

    def test_transition_row_reads_measured_value_and_result_without_shifting_marker(self):
        document = Document()
        table = document.add_table(rows=2, cols=8)
        values = ["D55", "卸油区-金属法兰-基准点", "", "铜编织带", "LPZ0B", "≤0.03", "0.02", "符合"]
        for column, value in enumerate(values):
            table.cell(1, column).text = value

        row = parse_transition_row(table, 1)

        self.assertEqual(row["marker"], "D55")
        self.assertEqual(row["measuredValue"], "0.02")
        self.assertEqual(row["result"], "符合")

    def test_grounding_single_location_cell_value_is_equipment_name(self):
        document = Document()
        table = document.add_table(rows=2, cols=8)
        values = ["1", "1#加油机", "", "—", "LPZ0B", "≤4.0", "1.2", "符合"]
        for column, value in enumerate(values):
            table.cell(1, column).text = value

        row = parse_grounding_row(table, 1)

        self.assertEqual(row["workLocation"], "")
        self.assertEqual(row["equipmentName"], "1#加油机")

    def test_transition_row_resolves_columns_when_standard_header_is_merged(self):
        document = Document()
        table = document.add_table(rows=3, cols=9)
        headers = ["编号", "所在位置、设备名称、基准点", "", "连接导体材质规格", "防雷分区", "标准值（Ω）", "", "测试值（Ω）", "结论"]
        for column, value in enumerate(headers):
            table.cell(1, column).text = value
        table.cell(1, 1).merge(table.cell(1, 2))
        table.cell(1, 5).merge(table.cell(1, 6))
        values = ["D1", "站内，金属管道法兰，基准点", "", "—", "LPZ1", "≤0.03", "", "0.010", "符合"]
        for column, value in enumerate(values):
            table.cell(2, column).text = value
        table.cell(2, 1).merge(table.cell(2, 2))
        table.cell(2, 5).merge(table.cell(2, 6))

        row = parse_transition_row(table, 2)

        self.assertEqual(row["standardValue"], "≤0.03")
        self.assertEqual(row["measuredValue"], "0.010")
        self.assertEqual(row["result"], "符合")

    def test_matches_each_report_row_to_one_cad_range_text(self):
        texts = [{
            "id": 7,
            "importedSourceHandles": ["RANGE-TEXT"],
            "text": "D55~D56",
            "x": 110,
            "y": 80,
        }]

        first = _match_marker_text("D55", {}, texts, [])
        second = _match_marker_text("D56", {}, texts, [])
        outside = _match_marker_text("D57", {}, texts, [])

        self.assertEqual(first["sourceElementIds"], [7])
        self.assertEqual(second["sourceElementIds"], [7])
        self.assertIsNone(outside)

    def test_matches_compact_range_with_omitted_end_prefix(self):
        texts = [{
            "id": 8,
            "importedSourceHandles": ["RANGE-D1-30"],
            "text": "D1-30",
            "x": 110,
            "y": 80,
        }]

        self.assertIsNotNone(_match_marker_text("D1", {}, texts, []))
        self.assertIsNotNone(_match_marker_text("D16", {}, texts, []))
        self.assertIsNotNone(_match_marker_text("D30", {}, texts, []))
        self.assertIsNone(_match_marker_text("D31", {}, texts, []))

    def test_matches_zero_padded_cad_grounding_markers_to_word_numbers(self):
        texts = [
            {
                "id": 9,
                "importedSourceHandles": ["GROUND-005"],
                "text": "005",
                "x": 110,
                "y": 80,
            },
            {
                "id": 10,
                "importedSourceHandles": ["TRANSITION-D005"],
                "text": "D005",
                "x": 150,
                "y": 80,
            },
        ]

        grounding = _match_marker_text("5", {}, texts, [])
        transition = _match_marker_text("D5", {}, texts, [])

        self.assertEqual(grounding["sourceElementIds"], [9])
        self.assertEqual(transition["sourceElementIds"], [10])

    def test_prefers_location_code_over_matching_report_ordinal(self):
        texts = [
            {
                "id": 11,
                "importedSourceHandles": ["DL-005"],
                "text": "005",
                "x": 110,
                "y": 80,
            },
            {
                "id": 12,
                "importedSourceHandles": ["DL-001"],
                "text": "001",
                "x": 150,
                "y": 80,
            },
        ]

        first = _match_marker_text(
            "1",
            {"workLocation": "编号DL-005"},
            texts,
            [],
            report_type="grounding",
        )
        fifth = _match_marker_text(
            "5",
            {"workLocation": "编号DL-001"},
            texts,
            [],
            report_type="grounding",
        )

        self.assertEqual(first["sourceElementIds"], [11])
        self.assertEqual(fifth["sourceElementIds"], [12])

    def test_prefers_plain_numbered_location_over_report_ordinal(self):
        texts = [
            {
                "id": 101,
                "importedSourceHandles": ["CAD-008"],
                "text": "008",
                "x": 110,
                "y": 80,
            },
            {
                "id": 102,
                "importedSourceHandles": ["CAD-010"],
                "text": "010",
                "x": 150,
                "y": 80,
            },
        ]

        target = _match_marker_text(
            "8",
            {"workLocation": "编号010，动力接线箱"},
            texts,
            [],
            report_type="grounding",
        )

        self.assertEqual(target["sourceElementIds"], [102])

    def test_plain_numbered_locations_keep_zero_padding_and_map_independently(self):
        texts = [
            {
                "id": 201,
                "importedSourceHandles": ["CAD-008"],
                "text": "008",
                "x": 110,
                "y": 80,
            },
            {
                "id": 202,
                "importedSourceHandles": ["CAD-010"],
                "text": "010",
                "x": 150,
                "y": 80,
            },
            {
                "id": 203,
                "importedSourceHandles": ["CAD-012"],
                "text": "012",
                "x": 190,
                "y": 80,
            },
        ]

        cases = [
            ("7", "编号008，机柜间", 201),
            ("8", "编号010，动力接线箱", 202),
            ("13", "编号012，太阳能光伏板支架", 203),
        ]
        for report_marker, work_location, expected_id in cases:
            with self.subTest(report_marker=report_marker):
                target = _match_marker_text(
                    report_marker,
                    {"workLocation": work_location},
                    texts,
                    [],
                    report_type="grounding",
                )
                self.assertEqual(target["sourceElementIds"], [expected_id])

    def test_does_not_fall_back_to_report_ordinal_when_explicit_code_is_missing(self):
        texts = [
            {
                "id": 301,
                "importedSourceHandles": ["CAD-033"],
                "text": "033",
                "x": 110,
                "y": 80,
            },
        ]

        target = _match_marker_text(
            "33",
            {"workLocation": "编号026，缺少的图纸测试点"},
            texts,
            [],
            report_type="grounding",
        )

        self.assertIsNone(target)

    def test_location_code_matching_is_limited_to_grounding_rows(self):
        texts = [
            {
                "id": 15,
                "importedSourceHandles": ["DL-005"],
                "text": "005",
                "x": 110,
                "y": 80,
            },
            {
                "id": 16,
                "importedSourceHandles": ["TRANSITION-D1"],
                "text": "D1",
                "x": 150,
                "y": 80,
            },
        ]

        transition = _match_marker_text(
            "D1",
            {"workLocation": "编号DL-005"},
            texts,
            [],
            report_type="transition",
        )

        self.assertEqual(transition["sourceElementIds"], [16])

    def test_does_not_treat_equipment_model_as_grounding_location_code(self):
        texts = [
            {
                "id": 17,
                "importedSourceHandles": ["EQUIPMENT-BV102"],
                "text": "102",
                "x": 110,
                "y": 80,
            },
            {
                "id": 18,
                "importedSourceHandles": ["REPORT-16"],
                "text": "16",
                "x": 150,
                "y": 80,
            },
        ]

        target = _match_marker_text(
            "16",
            {"workLocation": "无编号，BV102接地点"},
            texts,
            [],
            report_type="grounding",
        )

        self.assertEqual(target["sourceElementIds"], [18])

    def test_keeps_zero_padded_location_codes_distinct_from_report_numbers(self):
        texts = [
            {
                "id": 13,
                "importedSourceHandles": ["DL-034"],
                "text": "03\n4",
                "x": 110,
                "y": 80,
            },
            {
                "id": 14,
                "importedSourceHandles": ["REPORT-34"],
                "text": "34",
                "x": 150,
                "y": 80,
            },
        ]

        location_bound = _match_marker_text(
            "37",
            {"workLocation": "编号DL-034"},
            texts,
            [],
            report_type="grounding",
        )
        ordinal_bound = _match_marker_text(
            "34",
            {"workLocation": "无编号"},
            texts,
            [],
            report_type="grounding",
        )

        self.assertEqual(location_bound["sourceElementIds"], [13])
        self.assertEqual(ordinal_bound["sourceElementIds"], [14])

    def test_dl_034_keeps_its_current_report_row_instead_of_row_34(self):
        report_tables = {
            "grounding": [
                {"marker": "34", "workLocation": "无编号，放空区静电泄放装置接地点"},
                {"marker": "37", "workLocation": "编号DL-034，新建监控摄像头接地点"},
            ],
        }

        marker, replacement = _report_marker_for_display_label(
            "034",
            "37",
            "grounding",
            report_tables,
        )

        self.assertEqual(marker, "37")
        self.assertIsNone(replacement)

    def test_displays_word_marker_and_retains_cad_location_marker_as_source(self):
        document = ezdxf.new("R2018")
        document.modelspace().add_text("005", dxfattribs={"height": 2.5}).set_placement((20, 30))
        stream = StringIO()
        document.write(stream)
        binding_data = {
            "reportTables": {
                "grounding": [{
                    "marker": "1",
                    "workLocation": "编号DL-005",
                    "equipmentName": "围墙接地点",
                }],
            },
        }

        result = parse_dwg_workspace(stream.getvalue().encode("utf-8"), "sample.dxf", binding_data)
        canvas = result["workspace"]["tabData"]["1"]
        point = canvas["testPoints"][0]

        self.assertEqual(point["label"], "1")
        self.assertEqual(point["reportMarker"], "1")
        self.assertEqual(point["sourceMarker"], "005")
        self.assertEqual(point["reportFields"]["workLocation"], "编号DL-005")
        self.assertEqual(point["reportFields"]["equipmentName"], "围墙接地点")

        finalize_existing_report_workspace(result["workspace"])
        source_label = next(text for text in canvas["texts"] if text.get("text") == "005")
        self.assertTrue(source_label["hidden"])

    def test_imports_cad_range_as_one_visible_point_with_two_report_rows(self):
        document = ezdxf.new("R2018")
        document.modelspace().add_text("D55~D56", dxfattribs={"height": 2.5}).set_placement((20, 30))
        stream = StringIO()
        document.write(stream)
        binding_data = {
            "reportTables": {
                "transition": [
                    {"marker": "D55", "equipmentName": "金属法兰", "measuredValue": "0.01"},
                    {"marker": "D56", "equipmentName": "金属法兰", "measuredValue": "0.02"},
                ],
            },
        }

        result = parse_dwg_workspace(stream.getvalue().encode("utf-8"), "range.dxf", binding_data)
        points = result["workspace"]["tabData"]["1"]["testPoints"]

        self.assertEqual(points[0]["label"], "D55-D56")
        self.assertEqual(points[0]["reportMarkers"], ["D55", "D56"])
        self.assertFalse(points[0]["reportOnly"])
        self.assertEqual(points[0]["reportFields"]["measuredValue"], "0.01")
        self.assertTrue(points[1]["reportOnly"])
        self.assertEqual(points[1]["reportMarker"], "D56")
        self.assertEqual(points[1]["reportFields"]["measuredValue"], "0.02")
        self.assertEqual(result["stats"]["boundRows"], 2)
        self.assertEqual(result["unmatched"], [])

    def test_collapses_consecutive_flange_rows_into_one_visible_test_point(self):
        points = [
            {
                "id": 10,
                "label": "D55",
                "reportFields": {"equipmentName": "金属法兰"},
                "reportType": "equipotentialBonding",
                "sourceElementIds": [1],
                "sourceHandles": ["TEXT-55"],
                "x": 100,
                "y": 80,
            },
            {
                "id": 11,
                "label": "D56",
                "reportFields": {"equipmentName": "金属法兰"},
                "reportType": "equipotentialBonding",
                "sourceElementIds": [2],
                "sourceHandles": ["TEXT-56"],
                "x": 120,
                "y": 80,
            },
        ]
        paths = [{
            "closed": True,
            "id": 3,
            "importedSourceHandles": ["FLANGE"],
            "points": [
                {"x": 100, "y": 90},
                {"x": 120, "y": 90},
                {"x": 120, "y": 100},
                {"x": 100, "y": 100},
            ],
        }]

        result = _collapse_paired_flange_test_points(points, paths)

        self.assertEqual(result[0]["label"], "D55~D56")
        self.assertEqual(result[0]["reportMarkers"], ["D55", "D56"])
        self.assertEqual(result[0]["reportMarker"], "D55")
        self.assertFalse(result[0]["reportOnly"])
        self.assertEqual(result[0]["sourceElementIds"], [1, 2, 3])
        self.assertTrue(result[1]["reportOnly"])
        self.assertEqual(result[1]["reportMarker"], "D56")
        self.assertEqual(result[1]["visualTestPointId"], 10)

    def test_keeps_consecutive_grounding_rows_as_independent_points(self):
        points = [
            {
                "id": 20,
                "label": "6",
                "reportFields": {"equipmentName": "接地点"},
                "reportType": "grounding",
                "sourceElementIds": [1],
                "x": 100,
                "y": 80,
            },
            {
                "id": 21,
                "label": "7",
                "reportFields": {"equipmentName": "接地点"},
                "reportType": "grounding",
                "sourceElementIds": [1],
                "x": 120,
                "y": 80,
            },
        ]

        result = _collapse_paired_flange_test_points(points, [])

        self.assertEqual([point["label"] for point in result], ["6", "7"])
        self.assertFalse(any(point.get("reportOnly") for point in result))

    def test_removes_empty_frame_left_of_metal_roof_label(self):
        empty_frame = {
            "closed": True,
            "id": 1,
            "name": "HATCH",
            "points": [
                {"x": 100, "y": 100},
                {"x": 200, "y": 100},
                {"x": 200, "y": 145},
                {"x": 100, "y": 145},
            ],
        }
        actual_hatch = {
            "closed": True,
            "fillStyle": "hatch",
            "id": 2,
            "name": "HATCH",
            "points": [
                {"x": 110, "y": 210},
                {"x": 210, "y": 210},
                {"x": 210, "y": 255},
                {"x": 110, "y": 255},
            ],
        }
        texts = [{
            "height": 32,
            "name": "\u91d1\u5c5e\u5c4b\u9762",
            "text": "\u91d1\u5c5e\u5c4b\u9762",
            "width": 120,
            "x": 220,
            "y": 106,
        }]

        result = _remove_metal_roof_legend_empty_frame(
            [empty_frame, actual_hatch],
            texts,
            {"x": 80, "y": 80, "width": 300, "height": 220},
        )

        self.assertEqual([path["id"] for path in result], [2])

    def test_keeps_metal_roof_hatch_and_removes_coincident_empty_outline(self):
        actual_hatch = {
            "closed": True,
            "fillStyle": "hatch",
            "hatch": True,
            "id": 1,
            "name": "HATCH",
            "points": [
                {"x": 100, "y": 100},
                {"x": 170, "y": 100},
                {"x": 170, "y": 130},
                {"x": 100, "y": 130},
            ],
        }
        empty_outline = {
            "closed": True,
            "id": 2,
            "name": "LWPOLYLINE",
            "points": actual_hatch["points"],
        }
        texts = [{
            "height": 24,
            "name": "\u91d1\u5c5e\u5c4b\u9762",
            "text": "\u91d1\u5c5e\u5c4b\u9762",
            "width": 90,
            "x": 180,
            "y": 103,
        }]

        result = _remove_metal_roof_legend_empty_frame(
            [actual_hatch, empty_outline],
            texts,
            {"x": 80, "y": 80, "width": 240, "height": 100},
        )

        self.assertEqual([path["id"] for path in result], [1])

    def test_removes_dash_title_swatch_after_default_dwg_normalization(self):
        swatch = {
            "closed": True,
            "id": 1,
            "name": "HATCH",
            "patternName": "DASH",
            "points": [
                {"x": 1213.25, "y": 920.27},
                {"x": 1271.01, "y": 920.27},
                {"x": 1271.01, "y": 924.24},
                {"x": 1213.25, "y": 924.24},
            ],
        }

        self.assertTrue(_is_orphan_legend_swatch(swatch, 1600, 1280))

    def test_removes_upper_stacked_legend_hatch_frame(self):
        upper = {
            "closed": True,
            "id": 1,
            "importedSourceHandles": ["HATCH-1"],
            "name": "HATCH",
            "points": [
                {"x": 100, "y": 100},
                {"x": 160, "y": 100},
                {"x": 160, "y": 130},
                {"x": 100, "y": 130},
            ],
        }
        lower = {
            "closed": True,
            "fillStyle": "hatch",
            "id": 2,
            "importedSourceHandles": ["HATCH-1"],
            "name": "HATCH",
            "points": [
                {"x": 100, "y": 140},
                {"x": 160, "y": 140},
                {"x": 160, "y": 170},
                {"x": 100, "y": 170},
            ],
        }

        result = _merge_legend_hatch_parts(
            [upper, lower],
            {"x": 80, "y": 80, "width": 120, "height": 120},
        )

        self.assertEqual([path["id"] for path in result], [1])
        self.assertEqual(result[0]["fillStyle"], "solid")

    def test_merges_separately_handled_halves_of_legend_triangle(self):
        upper = {
            "closed": True,
            "fillStyle": "solid",
            "id": 1,
            "importedSourceHandles": ["HATCH-UPPER"],
            "name": "HATCH",
            "points": [
                {"x": 10, "y": 0}, {"x": 0, "y": 5},
                {"x": 10, "y": 5}, {"x": 10, "y": 0},
            ],
        }
        lower = {
            "closed": True,
            "fillStyle": "solid",
            "id": 2,
            "importedSourceHandles": ["HATCH-LOWER"],
            "name": "HATCH",
            "points": [
                {"x": 10, "y": 5}, {"x": 0, "y": 5},
                {"x": 10, "y": 10}, {"x": 10, "y": 5},
            ],
        }

        result = _merge_legend_hatch_parts(
            [upper, lower],
            {"x": 0, "y": 0, "width": 40, "height": 40},
        )

        self.assertEqual([path["id"] for path in result], [1])
        self.assertEqual(_triangle_side_from_path(result[0]), "left")

    def test_removes_only_large_polyline_outline_overlapping_hatch(self):
        hatch = {
            "closed": True,
            "fillStyle": "hatch",
            "id": 1,
            "name": "HATCH",
            "points": [
                {"x": 100, "y": 100},
                {"x": 500, "y": 100},
                {"x": 500, "y": 350},
                {"x": 100, "y": 350},
            ],
        }
        duplicate_outline = {
            "closed": True,
            "id": 2,
            "name": "LWPOLYLINE",
            "points": [
                {"x": 100, "y": 150},
                {"x": 500, "y": 150},
                {"x": 500, "y": 400},
                {"x": 100, "y": 400},
            ],
        }
        unrelated_small_outline = {
            "closed": True,
            "id": 3,
            "name": "LWPOLYLINE",
            "points": [
                {"x": 20, "y": 20},
                {"x": 40, "y": 20},
                {"x": 40, "y": 40},
                {"x": 20, "y": 40},
            ],
        }

        result = _remove_redundant_large_hatch_outlines(
            [hatch, duplicate_outline, unrelated_small_outline],
            1600,
            1280,
        )

        self.assertEqual([path["id"] for path in result], [1, 3])

    def test_marker_interaction_group_prefers_exact_element_ids(self):
        parsed = {
            "blocks": [],
            "paths": [
                {"id": 1, "importedSourceHandles": ["SHARED"]},
                {"id": 2, "importedSourceHandles": ["SHARED"]},
            ],
            "texts": [{"id": 3, "importedSourceHandles": ["LABEL"]}],
        }
        point = {
            "id": 4,
            "label": "D1",
            "sourceElementIds": [1, 3],
            "sourceHandles": ["SHARED", "LABEL"],
        }

        _assign_interaction_groups(parsed, [point])

        self.assertEqual(parsed["paths"][0]["interactionGroupId"], "cad-marker-D1")
        self.assertNotIn("interactionGroupId", parsed["paths"][1])
        self.assertEqual(parsed["texts"][0]["interactionGroupId"], "cad-marker-D1")
        self.assertEqual(point["interactionGroupId"], "cad-marker-D1")

    def test_existing_report_marker_keeps_editable_interaction_group(self):
        workspace = {
            "activeTabId": 1,
            "tabData": {
                "1": {
                    "boardWidth": 1000,
                    "boardHeight": 800,
                    "nextId": 4,
                    "paths": [{
                        "closed": True,
                        "id": 1,
                        "interactionGroupId": "cad-marker-D1",
                        "name": "SOLID",
                        "points": [
                            {"x": 390, "y": 390},
                            {"x": 410, "y": 390},
                            {"x": 410, "y": 410},
                            {"x": 390, "y": 410},
                        ],
                    }],
                    "reportTables": {},
                    "testPoints": [{
                        "id": 3,
                        "interactionGroupId": "cad-marker-D1",
                        "label": "D1",
                        "side": "right",
                        "size": 0.38,
                        "x": 400,
                        "y": 400,
                    }],
                    "texts": [],
                },
            },
        }

        finalize_existing_report_workspace(workspace)
        canvas = workspace["tabData"]["1"]

        self.assertEqual(canvas["testPoints"][0]["interactionGroupId"], "cad-marker-D1")
        self.assertTrue(canvas["testPoints"][0]["cadSourceVisible"])
        self.assertFalse(canvas["paths"][0].get("locked", False))
        self.assertTrue(canvas["paths"][0]["testPointSource"])

    def test_uses_field_labels_for_blank_report_values_except_marker(self):
        fields = _report_fields(
            {
                "marker": "D1",
                "equipmentName": "金属管道法兰",
                "referencePoint": None,
                "measuredValue": "—",
            },
            "transition",
        )

        self.assertNotIn("marker", fields)
        self.assertEqual(fields["equipmentName"], "金属管道法兰")
        self.assertEqual(fields["referencePoint"], "基准点")
        self.assertEqual(fields["conductorSpec"], "连接导体材质规格")
        self.assertEqual(fields["measuredValue"], "—")

    def test_fits_outlined_equipment_text_inside_enclosing_box(self):
        texts = [{
            "fontSize": 20,
            "glyphBounds": {"x": 0, "y": 0, "width": 80, "height": 30},
            "glyphPaths": [[
                {"x": 0, "y": 0}, {"x": 80, "y": 0},
                {"x": 80, "y": 30}, {"x": 0, "y": 30},
            ]],
            "rotation": 0,
            "text": "过滤器",
            "x": 10,
            "y": 10,
        }]
        paths = [{
            "closed": True,
            "points": [
                {"x": 8, "y": 8}, {"x": 68, "y": 8},
                {"x": 68, "y": 42}, {"x": 8, "y": 42},
            ],
        }]

        _fit_cad_text_to_enclosing_boxes(texts, paths)

        bounds = texts[0]["glyphBounds"]
        self.assertLessEqual(bounds["width"], 60 * 0.86 + 0.001)
        self.assertLessEqual(bounds["height"], 34 * 0.82 + 0.001)
        self.assertAlmostEqual(texts[0]["x"] + bounds["x"] + bounds["width"] / 2, 38)

    def _dxf_bytes(self):
        document = ezdxf.new("R2018")
        modelspace = document.modelspace()
        modelspace.add_line((0, 0), (100, 100))
        modelspace.add_text("D3", dxfattribs={"height": 2.5}).set_placement((20, 30))
        stream = StringIO()
        document.write(stream)
        return stream.getvalue().encode("utf-8")

    def _numbered_dxf_bytes(self):
        document = ezdxf.new("R2018")
        modelspace = document.modelspace()
        for index, label in enumerate(("1", "2", "3", "FL-1", "TT1101")):
            modelspace.add_text(label, dxfattribs={"height": 2.5}).set_placement((index * 10, 10))
        stream = StringIO()
        document.write(stream)
        return stream.getvalue().encode("utf-8")

    def test_binds_report_row_to_matching_drawing_label(self):
        binding_data = {
            "reportTables": {
                "grounding": [
                    {
                        "id": 50,
                        "marker": "D3",
                        "workLocation": "静电泄放装置",
                        "measuredValue": "1.44",
                    }
                ]
            }
        }

        result = parse_dwg_workspace(self._dxf_bytes(), "sample.dxf", binding_data)
        canvas = result["workspace"]["tabData"]["1"]
        point = canvas["testPoints"][0]
        row = canvas["reportTables"]["grounding"][0]

        self.assertEqual(point["id"], row["id"])
        self.assertEqual(point["label"], "D3")
        self.assertEqual(point["reportFields"]["measuredValue"], "1.44")
        self.assertEqual(result["stats"]["boundRows"], 1)
        self.assertEqual(result["unmatched"], [])

    def test_reports_unmatched_rows_without_creating_fake_points(self):
        binding_data = {
            "reportTables": {
                "grounding": [{"marker": "D9", "workLocation": "不存在的点"}]
            }
        }

        result = parse_dwg_workspace(self._dxf_bytes(), "sample.dxf", binding_data)
        canvas = result["workspace"]["tabData"]["1"]

        self.assertEqual(canvas["testPoints"], [])
        self.assertEqual(result["stats"]["unmatchedRows"], 1)

    def test_d_prefixed_report_rows_do_not_replace_numeric_drawing_points(self):
        document = ezdxf.new("R2018")
        modelspace = document.modelspace()
        for index in range(1, 4):
            modelspace.add_text(str(index), dxfattribs={"height": 2.5}).set_placement((index * 10, 10))
        stream = StringIO()
        document.write(stream)
        binding_data = {
            "reportTables": {
                "grounding": [
                    {"marker": str(index), "measuredValue": f"1.{index}"}
                    for index in range(1, 4)
                ],
                "transition": [
                    {"marker": f"D{index}", "measuredValue": f"0.0{index}"}
                    for index in range(1, 4)
                ],
            }
        }

        result = parse_dwg_workspace(stream.getvalue().encode("utf-8"), "sample.dxf", binding_data)
        canvas = result["workspace"]["tabData"]["1"]

        self.assertEqual([point["label"] for point in canvas["testPoints"]], ["1", "2", "3"])
        self.assertEqual([point["reportType"] for point in canvas["testPoints"]], ["grounding"] * 3)
        self.assertEqual(result["stats"]["boundRows"], 3)
        self.assertEqual(
            [item["marker"] for item in result["unmatched"]],
            ["D1", "D2", "D3"],
        )

    def test_detects_continuous_numeric_points_without_binding_data(self):
        result = parse_dwg_workspace(self._numbered_dxf_bytes(), "sample.dxf")
        points = result["workspace"]["tabData"]["1"]["testPoints"]

        self.assertEqual([point["label"] for point in points], ["1", "2", "3"])

    def test_declared_transition_range_creates_one_special_visible_point(self):
        document = ezdxf.new("R2018")
        modelspace = document.modelspace()
        for index, label in enumerate(("1", "2", "3"), start=1):
            modelspace.add_text(label, dxfattribs={"height": 2.5}).set_placement((index * 10, 10))
        modelspace.add_text("D1-D3", dxfattribs={"height": 2.5}).set_placement((10, 30))
        stream = StringIO()
        document.write(stream)
        binding_data = {
            "reportTables": {
                "transition": [
                    {"marker": "", "measuredValue": "0.01"},
                    {"marker": "", "measuredValue": "0.02"},
                    {"marker": "", "measuredValue": "0.03"},
                ]
            }
        }

        result = parse_dwg_workspace(stream.getvalue().encode("utf-8"), "sample.dxf", binding_data)
        canvas = result["workspace"]["tabData"]["1"]

        self.assertEqual([point["label"] for point in canvas["testPoints"]], ["D1-D3", "D2", "D3"])
        self.assertEqual(canvas["testPoints"][0]["reportMarkers"], ["D1", "D2", "D3"])
        self.assertTrue(canvas["testPoints"][0]["specialRange"])
        self.assertTrue(canvas["testPoints"][1]["reportOnly"])
        self.assertEqual(
            [row["marker"] for row in canvas["reportTables"]["transition"]],
            ["D1", "D2", "D3"],
        )
        self.assertEqual(result["stats"]["boundRows"], 3)
        self.assertEqual(result["unmatched"], [])

    def test_preserves_existing_workspace_point_position_and_binding(self):
        binding_data = {
            "boardWidth": 800,
            "boardHeight": 640,
            "testPoints": [
                {
                    "id": 50,
                    "label": "D3",
                    "x": 200,
                    "y": 320,
                    "side": "left",
                    "size": 0.45,
                    "sourceHandles": ["OLD"],
                    "binding": {"id": 9, "kind": "path"},
                }
            ],
            "reportTables": {
                "grounding": [{"id": 50, "marker": "D3", "measuredValue": "1.44"}]
            },
        }

        result = parse_dwg_workspace(
            self._dxf_bytes(),
            "sample.dxf",
            binding_data,
            board_width=1600,
            board_height=1280,
        )
        point = result["workspace"]["tabData"]["1"]["testPoints"][0]

        self.assertEqual((point["x"], point["y"]), (400, 640))
        self.assertEqual(point["side"], "left")
        self.assertEqual(point["size"], 0.45)
        self.assertEqual(point["binding"], {"id": 9, "kind": "path"})

    def test_reimport_replaces_stale_ordinal_binding_with_explicit_physical_code(self):
        document = ezdxf.new("R2018")
        modelspace = document.modelspace()
        modelspace.add_text("008", dxfattribs={"height": 2.5}).set_placement((20, 30))
        modelspace.add_text("010", dxfattribs={"height": 2.5}).set_placement((80, 30))
        stream = StringIO()
        document.write(stream)
        binding_data = {
            "boardWidth": 800,
            "boardHeight": 640,
            "testPoints": [
                {
                    "id": 50,
                    "label": "8",
                    "reportMarker": "8",
                    "sourceMarker": "008",
                    "x": 200,
                    "y": 320,
                    "side": "left",
                },
            ],
            "reportTables": {
                "grounding": [
                    {
                        "id": 50,
                        "marker": "8",
                        "workLocation": "编号010，动力接线箱",
                    },
                ],
            },
        }

        result = parse_dwg_workspace(
            stream.getvalue().encode("utf-8"),
            "sample.dxf",
            binding_data,
            board_width=1600,
            board_height=1280,
        )
        point = result["workspace"]["tabData"]["1"]["testPoints"][0]

        self.assertEqual(point["reportMarker"], "8")
        self.assertEqual(point["sourceMarker"], "010")
        self.assertNotEqual((point["x"], point["y"]), (400, 640))

    def test_preserves_cad_fill_and_text_render_metadata(self):
        document = ezdxf.new("R2018")
        modelspace = document.modelspace()
        modelspace.add_solid([(0, 0), (10, 0), (0, 10), (0, 10)])
        modelspace.add_text("图例", dxfattribs={"height": 3}).set_placement((20, 20))
        stream = StringIO()
        document.write(stream)

        result = parse_dwg_workspace(stream.getvalue().encode("utf-8"), "sample.dxf")
        canvas = result["workspace"]["tabData"]["1"]
        solid = next(path for path in canvas["paths"] if path["name"] == "SOLID")
        text = next(item for item in canvas["texts"] if item["text"] == "图例")

        self.assertEqual(solid["fillStyle"], "solid")
        self.assertEqual(solid["fillColor"], "#111111")
        self.assertTrue(solid["cadRender"])
        self.assertTrue(text["cadRender"])
        self.assertIn("fontFamily", text)
        self.assertTrue(text["glyphPaths"])

    def test_preserves_inline_font_numeric_mtext_as_cad_glyphs(self):
        document = ezdxf.new("R2018")
        document.modelspace().add_mtext(
            r"{\fArial|b0|i0|c0|p34;20}",
            dxfattribs={"char_height": 2.5, "attachment_point": 5},
        ).set_location((20, 30))
        stream = StringIO()
        document.write(stream)

        result = parse_dwg_workspace(stream.getvalue().encode("utf-8"), "number.dxf")
        text = result["workspace"]["tabData"]["1"]["texts"][0]

        self.assertEqual(text["text"], "20")
        self.assertTrue(text["glyphPaths"])
        self.assertTrue(text["glyphBounds"]["width"] > 0)
        self.assertTrue(text["glyphBounds"]["height"] > 0)

    def test_preserves_hatch_pattern_and_point_marker(self):
        document = ezdxf.new("R2018")
        modelspace = document.modelspace()
        hatch = modelspace.add_hatch(color=7)
        hatch.set_pattern_fill("STEEL", scale=1.2, angle=0)
        hatch.paths.add_polyline_path(
            [(0, 0), (20, 0), (20, 8), (0, 8)],
            is_closed=True,
        )
        modelspace.add_point((10, 4))
        stream = StringIO()
        document.write(stream)

        result = parse_dwg_workspace(stream.getvalue().encode("utf-8"), "sample.dxf")
        paths = result["workspace"]["tabData"]["1"]["paths"]
        patterned = next(path for path in paths if path.get("patternName") == "STEEL")
        marker = next(path for path in paths if path.get("name") == "POINT")

        self.assertEqual(patterned["fillStyle"], "hatch")
        self.assertEqual(patterned["hatchAngle"], 45.0)
        self.assertEqual(patterned["hatchAngles"], [45.0])
        self.assertGreaterEqual(patterned["hatchSpacing"], 3)
        self.assertEqual(marker["fillStyle"], "solid")
        self.assertEqual(marker["fillColor"], "#111111")
        self.assertTrue(marker["closed"])

    def test_preserves_bold_title_block_wrap(self):
        document = ezdxf.new("R2018")
        modelspace = document.modelspace()
        value = "国家管网集团北方管道有限责任公司天津输油气分公司"
        modelspace.add_mtext(
            r"{\fSimSun|b1|i0|c134|p2;" + value + "}",
            dxfattribs={"char_height": 4, "width": 69.1519, "attachment_point": 5},
        ).set_location((10, 10))
        stream = StringIO()
        document.write(stream)

        result = parse_dwg_workspace(stream.getvalue().encode("utf-8"), "sample.dxf")
        text = result["workspace"]["tabData"]["1"]["texts"][0]

        self.assertEqual([len(line) for line in text["text"].split("\n")], [12, 12])
        self.assertEqual(text["fontWeight"], 700)
        self.assertTrue(text["glyphPaths"])

    def test_balances_short_mtext_into_two_character_lines(self):
        document = ezdxf.new("R2018")
        modelspace = document.modelspace()
        modelspace.add_mtext(
            "放空立管",
            dxfattribs={"char_height": 5, "width": 7.6},
        ).set_location((10, 10))
        stream = StringIO()
        document.write(stream)

        result = parse_dwg_workspace(stream.getvalue().encode("utf-8"), "sample.dxf")
        text = result["workspace"]["tabData"]["1"]["texts"][0]

        self.assertEqual(text["text"], "放空\n立管")
